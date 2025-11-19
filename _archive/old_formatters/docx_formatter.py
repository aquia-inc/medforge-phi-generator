"""
DOCX document formatter for PHI documents
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os


class PHIDocxFormatter:
    """Creates DOCX documents with PHI content"""

    def __init__(self, output_dir='output'):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def create_lab_result(self, patient, provider, facility, lab_data, filename):
        """Generate a lab result document (PHI Positive)"""
        doc = Document()

        # Add facility header
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header.add_run(facility['name'].upper())
        run.bold = True
        run.font.size = Pt(16)

        # Facility address
        addr_para = doc.add_paragraph()
        addr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        addr_para.add_run(
            f"{facility['address']['street']}\n"
            f"{facility['address']['city']}, {facility['address']['state']} {facility['address']['zip']}\n"
            f"Phone: {facility['phone']} | Fax: {facility['fax']}"
        ).font.size = Pt(10)

        doc.add_paragraph()  # Spacing

        # Document title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run('LABORATORY RESULTS')
        title_run.bold = True
        title_run.font.size = Pt(14)

        doc.add_paragraph()  # Spacing

        # Patient Information Section
        doc.add_paragraph().add_run('PATIENT INFORMATION').bold = True
        patient_info = doc.add_table(rows=6, cols=2)
        patient_info.style = 'Light Grid Accent 1'

        # Fill patient info
        cells = [
            ('Patient Name:', f"{patient['last_name']}, {patient['first_name']}"),
            ('Date of Birth:', patient['dob'].strftime('%m/%d/%Y')),
            ('Age:', str(patient['age'])),
            ('MRN:', patient['mrn']),
            ('Address:', f"{patient['address']['street']}, {patient['address']['city']}, {patient['address']['state']} {patient['address']['zip']}"),
            ('Phone:', patient['phone'])
        ]

        for i, (label, value) in enumerate(cells):
            patient_info.rows[i].cells[0].text = label
            patient_info.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            patient_info.rows[i].cells[1].text = value

        doc.add_paragraph()  # Spacing

        # Test Information
        doc.add_paragraph().add_run('TEST INFORMATION').bold = True
        test_info = doc.add_table(rows=3, cols=2)
        test_info.style = 'Light Grid Accent 1'

        test_cells = [
            ('Collection Date:', lab_data['test_date'].strftime('%m/%d/%Y')),
            ('Report Date:', datetime.now().strftime('%m/%d/%Y')),
            ('Ordering Provider:', f"{provider['first_name']} {provider['last_name']}, {provider['title']}")
        ]

        for i, (label, value) in enumerate(test_cells):
            test_info.rows[i].cells[0].text = label
            test_info.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            test_info.rows[i].cells[1].text = value

        doc.add_paragraph()  # Spacing

        # Lab Results Table
        doc.add_paragraph().add_run('LABORATORY RESULTS').bold = True
        results_table = doc.add_table(rows=len(lab_data['results']) + 1, cols=5)
        results_table.style = 'Light Grid Accent 1'

        # Headers
        headers = ['Test Name', 'Result', 'Unit', 'Reference Range', 'Flag']
        for i, header in enumerate(headers):
            cell = results_table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True

        # Fill results
        for i, result in enumerate(lab_data['results'], 1):
            results_table.rows[i].cells[0].text = result['test']
            results_table.rows[i].cells[1].text = str(result['value'])
            results_table.rows[i].cells[2].text = result['unit']
            results_table.rows[i].cells[3].text = result['reference_range']
            flag_cell = results_table.rows[i].cells[4]
            flag_cell.text = result.get('flag', '')
            if result.get('flag'):
                # Highlight abnormal results
                flag_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
                flag_cell.paragraphs[0].runs[0].bold = True

        doc.add_paragraph()  # Spacing

        # Footer/Disclaimer
        footer = doc.add_paragraph()
        footer.add_run(
            '---\n'
            'This report contains confidential patient health information. '
            'Distribution or copying is prohibited without authorization.\n'
            f'Medical Director: {provider["first_name"]} {provider["last_name"]}, {provider["title"]}\n'
            f'NPI: {provider["npi"]}'
        ).font.size = Pt(8)

        # Save document
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        return filepath

    def create_progress_note(self, patient, provider, facility, filename):
        """Generate a clinical progress note (PHI Positive)"""
        doc = Document()

        # Header
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header.add_run(facility['name'].upper())
        run.bold = True
        run.font.size = Pt(14)

        doc.add_paragraph()

        # Document title
        title = doc.add_paragraph()
        title_run = title.add_run('PROGRESS NOTE / SOAP NOTE')
        title_run.bold = True
        title_run.font.size = Pt(12)

        doc.add_paragraph()

        # Patient header
        patient_header = doc.add_paragraph()
        patient_header.add_run(
            f"Patient: {patient['last_name']}, {patient['first_name']}  |  "
            f"DOB: {patient['dob'].strftime('%m/%d/%Y')}  |  "
            f"MRN: {patient['mrn']}\n"
            f"Date of Visit: {datetime.now().strftime('%m/%d/%Y')}\n"
            f"Provider: {provider['first_name']} {provider['last_name']}, {provider['title']} - {provider['specialty']}"
        ).font.size = Pt(10)

        doc.add_paragraph()

        # Vital Signs
        doc.add_paragraph().add_run('VITAL SIGNS').bold = True
        vitals = patient['vital_signs']
        vitals_text = (
            f"BP: {vitals['blood_pressure']} mmHg  |  "
            f"HR: {vitals['heart_rate']} bpm  |  "
            f"Temp: {vitals['temperature']}°F  |  "
            f"RR: {vitals['respiratory_rate']}  |  "
            f"O2 Sat: {vitals['oxygen_saturation']}%\n"
            f"Weight: {vitals['weight']} lbs  |  "
            f"Height: {vitals['height']} inches"
        )
        doc.add_paragraph(vitals_text)

        doc.add_paragraph()

        # SOAP Format
        # Subjective
        doc.add_paragraph().add_run('SUBJECTIVE:').bold = True
        doc.add_paragraph(
            f"Patient presents for follow-up visit. Reports feeling generally well. "
            f"Current medications include: {', '.join(patient['medications'][:3])}. "
            f"Known allergies: {', '.join(patient['allergies'])}."
        )

        # Objective
        doc.add_paragraph().add_run('OBJECTIVE:').bold = True
        doc.add_paragraph(
            f"Physical exam reveals patient in no acute distress. "
            f"Vital signs as noted above. "
            f"Review of systems within normal limits for age."
        )

        # Assessment
        doc.add_paragraph().add_run('ASSESSMENT:').bold = True
        for diagnosis in patient['diagnoses']:
            doc.add_paragraph(
                f"• {diagnosis['name']} (ICD-10: {diagnosis['icd10']})",
                style='List Bullet'
            )

        # Plan
        doc.add_paragraph().add_run('PLAN:').bold = True
        doc.add_paragraph(
            "1. Continue current medications as prescribed\n"
            "2. Follow-up lab work in 3 months\n"
            "3. Return to clinic in 6 months or sooner if symptoms worsen\n"
            "4. Patient education provided regarding disease management"
        )

        doc.add_paragraph()

        # Signature
        sig = doc.add_paragraph()
        sig.add_run(
            f"\n---\n"
            f"Electronically signed by: {provider['first_name']} {provider['last_name']}, {provider['title']}\n"
            f"Date: {datetime.now().strftime('%m/%d/%Y %H:%M')}\n"
            f"NPI: {provider['npi']}"
        ).font.size = Pt(9)

        # Save
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        return filepath

    def create_generic_medical_policy(self, facility, filename):
        """Generate a generic medical policy document (PHI Negative - No Patient Data)"""
        doc = Document()

        # Header
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header.add_run(facility['name'].upper())
        run.bold = True
        run.font.size = Pt(16)

        doc.add_paragraph()

        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run('CLINICAL PRACTICE POLICY')
        title_run.bold = True
        title_run.font.size = Pt(14)

        doc.add_paragraph()

        # Policy metadata
        meta = doc.add_table(rows=4, cols=2)
        meta.style = 'Light Grid Accent 1'
        meta_cells = [
            ('Policy Number:', 'CPG-2024-001'),
            ('Effective Date:', '01/01/2024'),
            ('Review Date:', '01/01/2025'),
            ('Department:', 'Clinical Operations')
        ]

        for i, (label, value) in enumerate(meta_cells):
            meta.rows[i].cells[0].text = label
            meta.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            meta.rows[i].cells[1].text = value

        doc.add_paragraph()

        # Purpose
        doc.add_paragraph().add_run('PURPOSE:').bold = True
        doc.add_paragraph(
            'This policy establishes guidelines for clinical documentation standards '
            'to ensure quality patient care and regulatory compliance.'
        )

        # Scope
        doc.add_paragraph().add_run('SCOPE:').bold = True
        doc.add_paragraph(
            'This policy applies to all clinical staff, including physicians, nurses, '
            'and allied health professionals providing patient care services.'
        )

        # Policy
        doc.add_paragraph().add_run('POLICY:').bold = True
        doc.add_paragraph(
            '1. All clinical encounters must be documented within 24 hours\n'
            '2. Documentation must include patient assessment, diagnosis, and treatment plan\n'
            '3. All entries must be signed and dated by the responsible provider\n'
            '4. Abbreviations must conform to the approved abbreviation list\n'
            '5. Late entries must be clearly identified as such'
        )

        # Procedure
        doc.add_paragraph().add_run('PROCEDURE:').bold = True
        doc.add_paragraph(
            'A. Documentation Requirements\n'
            '   - Chief complaint\n'
            '   - History of present illness\n'
            '   - Review of systems\n'
            '   - Physical examination findings\n'
            '   - Assessment and diagnosis\n'
            '   - Treatment plan\n\n'
            'B. Quality Assurance\n'
            '   - Random chart audits conducted quarterly\n'
            '   - Feedback provided to clinical staff\n'
            '   - Corrective action plans for deficiencies'
        )

        doc.add_paragraph()

        # Footer
        footer = doc.add_paragraph()
        footer.add_run(
            f'---\n'
            f'{facility["name"]}\n'
            f'{facility["address"]["street"]}, {facility["address"]["city"]}, {facility["address"]["state"]}\n'
            f'Policy Review Committee\n'
            f'Approved: {datetime.now().strftime("%m/%d/%Y")}'
        ).font.size = Pt(8)

        # Save
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        return filepath

    def create_blank_form_template(self, facility, filename):
        """Generate a blank patient intake form (PHI Negative - Template Only)"""
        doc = Document()

        # Header
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header.add_run(facility['name'].upper())
        run.bold = True
        run.font.size = Pt(14)

        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.add_run('PATIENT REGISTRATION FORM').bold = True

        doc.add_paragraph()

        # Instructions
        doc.add_paragraph(
            'Please complete all sections of this form. Information provided will be kept confidential.'
        ).italic = True

        doc.add_paragraph()

        # Patient Information Section
        doc.add_paragraph().add_run('PATIENT INFORMATION').bold = True

        # Create form fields with blank lines
        form_fields = [
            'Last Name: _______________________________________  First Name: _______________________________________',
            'Date of Birth: ______ / ______ / ______  Age: ______  Sex: ☐ M  ☐ F  ☐ Other',
            'Social Security Number: ______ - ______ - ______',
            'Address: _________________________________________________________________________________________',
            'City: ___________________________________  State: ______  ZIP: ______________',
            'Phone (Home): (______) ______-________  Phone (Cell): (______) ______-________',
            'Email: ____________________________________________________________________________________________',
        ]

        for field in form_fields:
            doc.add_paragraph(field)

        doc.add_paragraph()

        # Insurance Information
        doc.add_paragraph().add_run('INSURANCE INFORMATION').bold = True
        insurance_fields = [
            'Primary Insurance: _________________________________________________________________________',
            'Policy Number: ____________________________________________  Group Number: _________________',
            'Policy Holder Name: _______________________________________________________________________',
            'Relationship to Patient: ___________________________________________________________________',
        ]

        for field in insurance_fields:
            doc.add_paragraph(field)

        doc.add_paragraph()

        # Emergency Contact
        doc.add_paragraph().add_run('EMERGENCY CONTACT').bold = True
        emergency_fields = [
            'Name: ___________________________________________________________________________________________',
            'Relationship: _____________________________________________________________________________________',
            'Phone: (______) ______-________',
        ]

        for field in emergency_fields:
            doc.add_paragraph(field)

        doc.add_paragraph()

        # Signature
        doc.add_paragraph()
        doc.add_paragraph(
            'Patient Signature: ________________________________________________  Date: ______ / ______ / ______'
        )

        doc.add_paragraph()

        # Footer
        footer = doc.add_paragraph()
        footer.add_run(
            f'---\n'
            f'{facility["name"]} | {facility["phone"]}\n'
            f'This form is for use by authorized personnel only.'
        ).font.size = Pt(8)

        # Save
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        return filepath
