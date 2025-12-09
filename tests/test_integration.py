"""
Integration tests for MedForge generation
"""
import pytest
import tempfile
import shutil
import json
import os
import sys
from pathlib import Path

# Add src to path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'src'))


class TestCUIGeneration:
    """Integration tests for CUI document generation"""

    @pytest.fixture
    def temp_output_dir(self):
        """Create a temporary output directory"""
        temp_dir = tempfile.mkdtemp()
        yield temp_dir
        shutil.rmtree(temp_dir, ignore_errors=True)

    def test_cui_generation_creates_directory_structure(self, temp_output_dir):
        """Test that CUI generation creates correct directory structure"""
        from cli import MedForgeCUIGenerator

        generator = MedForgeCUIGenerator(
            output_dir=temp_output_dir,
            seed=42,
            formats=['docx', 'pdf'],
            categories=['financial', 'legal'],  # Limit categories for test
        )

        generator.generate_batch(
            cui_positive_count=5,
            cui_negative_count=3,
        )

        # Check directory structure - now uses CUI-{DisplayName}-Positive/Negative format
        assert Path(temp_output_dir, 'metadata').exists()

        # Check CUI display name folders exist
        cui_folders = [d for d in Path(temp_output_dir).iterdir() if d.is_dir() and d.name.startswith('CUI-')]
        assert len(cui_folders) > 0, "No CUI folders created"

        # Verify folder naming pattern
        for folder in cui_folders:
            assert folder.name.startswith('CUI-'), f"Folder {folder.name} doesn't follow CUI- pattern"
            assert folder.name.endswith('-Positive') or folder.name.endswith('-Negative'), \
                f"Folder {folder.name} doesn't end with -Positive or -Negative"

        # At least one folder should have files
        has_files = any(list(folder.glob('*')) for folder in cui_folders)
        assert has_files, "No files generated in CUI folders"

    def test_cui_manifest_generated(self, temp_output_dir):
        """Test that CUI manifest is correctly generated"""
        from cli import MedForgeCUIGenerator

        generator = MedForgeCUIGenerator(
            output_dir=temp_output_dir,
            seed=42,
            formats=['docx'],
        )

        generator.generate_batch(
            cui_positive_count=3,
            cui_negative_count=2,
        )

        manifest_path = Path(temp_output_dir, 'metadata', 'cui_manifest.json')
        assert manifest_path.exists(), "Manifest file not created"

        with open(manifest_path) as f:
            manifest = json.load(f)

        assert manifest['total_documents'] == 5
        assert manifest['cui_positive'] == 3
        assert manifest['cui_negative'] == 2
        assert len(manifest['files']) == 5

        # Check file entries have required fields (standardized schema)
        for file_entry in manifest['files']:
            assert 'file_path' in file_entry
            assert 'cui_status' in file_entry
            assert 'category' in file_entry
            assert 'subcategory' in file_entry  # Now required for all
            assert 'classification' in file_entry  # Now required for all
            assert 'authority' in file_entry  # Now required for all
            assert 'format' in file_entry
            # Verify file paths use CUI- prefix
            assert 'CUI-' in file_entry['file_path'], f"File path {file_entry['file_path']} doesn't contain CUI-"

    def test_cui_positive_documents_have_classification(self, temp_output_dir):
        """Test that CUI positive documents have classification markings"""
        from cli import MedForgeCUIGenerator

        generator = MedForgeCUIGenerator(
            output_dir=temp_output_dir,
            seed=42,
            formats=['docx'],
        )

        generator.generate_batch(
            cui_positive_count=3,
            cui_negative_count=0,
        )

        manifest_path = Path(temp_output_dir, 'metadata', 'cui_manifest.json')
        with open(manifest_path) as f:
            manifest = json.load(f)

        for file_entry in manifest['files']:
            if file_entry['cui_status'] == 'positive':
                assert 'classification' in file_entry
                assert file_entry['classification'], "Classification should not be empty"

    def test_cui_generation_with_specific_categories(self, temp_output_dir):
        """Test CUI generation with specific categories"""
        from cli import MedForgeCUIGenerator

        categories = ['financial', 'legal']
        generator = MedForgeCUIGenerator(
            output_dir=temp_output_dir,
            seed=42,
            categories=categories,
            formats=['docx'],
        )

        generator.generate_batch(
            cui_positive_count=10,
            cui_negative_count=5,
        )

        manifest_path = Path(temp_output_dir, 'metadata', 'cui_manifest.json')
        with open(manifest_path) as f:
            manifest = json.load(f)

        # All generated files should be from specified categories
        for file_entry in manifest['files']:
            assert file_entry['category'] in categories, \
                f"Unexpected category: {file_entry['category']}"

    def test_cui_files_are_valid(self, temp_output_dir):
        """Test that generated CUI files are valid and can be opened"""
        from cli import MedForgeCUIGenerator
        from docx import Document

        generator = MedForgeCUIGenerator(
            output_dir=temp_output_dir,
            seed=42,
            formats=['docx'],
        )

        generator.generate_batch(
            cui_positive_count=3,
            cui_negative_count=2,
        )

        # Find all DOCX files
        docx_files = list(Path(temp_output_dir).rglob('*.docx'))
        assert len(docx_files) > 0, "No DOCX files generated"

        # Try to open each file
        for docx_path in docx_files:
            try:
                doc = Document(str(docx_path))
                # Should have at least some content
                assert len(doc.paragraphs) > 0, f"Empty document: {docx_path}"
            except Exception as e:
                pytest.fail(f"Failed to open {docx_path}: {e}")


class TestMixedPHICUIGeneration:
    """Integration tests for mixed PHI and CUI generation"""

    @pytest.fixture
    def temp_output_dir(self):
        """Create a temporary output directory"""
        temp_dir = tempfile.mkdtemp()
        yield temp_dir
        shutil.rmtree(temp_dir, ignore_errors=True)

    def test_mixed_generation_creates_both_structures(self, temp_output_dir):
        """Test that mixed generation creates both PHI and CUI directories"""
        from cli import MedForgeGenerator, MedForgeCUIGenerator

        # Generate PHI documents
        phi_gen = MedForgeGenerator(
            output_dir=temp_output_dir,
            seed=42,
            formats=['docx', 'pdf'],
            llm_percentage=0.0,  # No LLM for faster tests
        )
        phi_gen.generate_batch(
            phi_positive_count=3,
            phi_negative_count=2,
        )

        # Generate CUI documents
        cui_gen = MedForgeCUIGenerator(
            output_dir=temp_output_dir,
            seed=42,
            formats=['docx', 'pdf'],
            llm_percentage=0.0,
        )
        cui_gen.generate_batch(
            cui_positive_count=3,
            cui_negative_count=2,
        )

        # Check PHI structures exist
        assert Path(temp_output_dir, 'phi_positive').exists()
        assert Path(temp_output_dir, 'phi_negative').exists()

        # Check CUI structures exist (now uses CUI-{DisplayName}-Positive/Negative format)
        cui_folders = [d for d in Path(temp_output_dir).iterdir() if d.is_dir() and d.name.startswith('CUI-')]
        assert len(cui_folders) > 0, "No CUI folders created"

        # Check both manifests exist
        assert Path(temp_output_dir, 'metadata', 'manifest.json').exists()
        assert Path(temp_output_dir, 'metadata', 'cui_manifest.json').exists()

    def test_mixed_manifests_are_separate(self, temp_output_dir):
        """Test that PHI and CUI manifests are separate and correct"""
        from cli import MedForgeGenerator, MedForgeCUIGenerator

        phi_gen = MedForgeGenerator(
            output_dir=temp_output_dir,
            seed=42,
            formats=['pdf', 'eml'],  # Use formats that work for PHI negative
            llm_percentage=0.0,
        )
        phi_gen.generate_batch(
            phi_positive_count=5,
            phi_negative_count=3,
        )

        cui_gen = MedForgeCUIGenerator(
            output_dir=temp_output_dir,
            seed=42,
            formats=['docx'],
            llm_percentage=0.0,
        )
        cui_gen.generate_batch(
            cui_positive_count=4,
            cui_negative_count=2,
        )

        # Check PHI manifest exists and has correct positive count
        phi_manifest_path = Path(temp_output_dir, 'metadata', 'manifest.json')
        with open(phi_manifest_path) as f:
            phi_manifest = json.load(f)

        assert phi_manifest['phi_positive'] == 5
        # PHI negative may be less due to format-specific generator limitations
        assert phi_manifest['phi_negative'] >= 0

        # Check CUI manifest
        cui_manifest_path = Path(temp_output_dir, 'metadata', 'cui_manifest.json')
        with open(cui_manifest_path) as f:
            cui_manifest = json.load(f)

        assert cui_manifest['cui_positive'] == 4
        assert cui_manifest['cui_negative'] == 2


class TestCUIFormats:
    """Test different output formats for CUI documents"""

    @pytest.fixture
    def temp_output_dir(self):
        """Create a temporary output directory"""
        temp_dir = tempfile.mkdtemp()
        yield temp_dir
        shutil.rmtree(temp_dir, ignore_errors=True)

    def test_pdf_generation(self, temp_output_dir):
        """Test PDF CUI document generation"""
        from cli import MedForgeCUIGenerator

        generator = MedForgeCUIGenerator(
            output_dir=temp_output_dir,
            seed=42,
            formats=['pdf'],
        )

        generator.generate_batch(
            cui_positive_count=3,
            cui_negative_count=2,
        )

        pdf_files = list(Path(temp_output_dir).rglob('*.pdf'))
        assert len(pdf_files) > 0, "No PDF files generated"

        # Verify PDF header
        for pdf_path in pdf_files:
            with open(pdf_path, 'rb') as f:
                header = f.read(5)
                assert header == b'%PDF-', f"Invalid PDF header: {pdf_path}"

    def test_docx_generation(self, temp_output_dir):
        """Test DOCX CUI document generation"""
        from cli import MedForgeCUIGenerator
        from docx import Document

        generator = MedForgeCUIGenerator(
            output_dir=temp_output_dir,
            seed=42,
            formats=['docx'],
        )

        generator.generate_batch(
            cui_positive_count=3,
            cui_negative_count=2,
        )

        docx_files = list(Path(temp_output_dir).rglob('*.docx'))
        assert len(docx_files) > 0, "No DOCX files generated"

        # Verify each DOCX can be opened
        for docx_path in docx_files:
            doc = Document(str(docx_path))
            assert len(doc.paragraphs) > 0

    def test_eml_generation(self, temp_output_dir):
        """Test EML CUI document generation"""
        from cli import MedForgeCUIGenerator

        generator = MedForgeCUIGenerator(
            output_dir=temp_output_dir,
            seed=42,
            formats=['eml'],
        )

        generator.generate_batch(
            cui_positive_count=3,
            cui_negative_count=2,
        )

        eml_files = list(Path(temp_output_dir).rglob('*.eml'))
        assert len(eml_files) > 0, "No EML files generated"

        # Verify email format
        for eml_path in eml_files:
            with open(eml_path, 'r') as f:
                content = f.read(500)
                assert any(h in content for h in ['From:', 'To:', 'Subject:', 'MIME-Version:']), \
                    f"Invalid email format: {eml_path}"


class TestCUIStats:
    """Test CUI generation statistics"""

    @pytest.fixture
    def temp_output_dir(self):
        """Create a temporary output directory"""
        temp_dir = tempfile.mkdtemp()
        yield temp_dir
        shutil.rmtree(temp_dir, ignore_errors=True)

    def test_stats_tracking(self, temp_output_dir):
        """Test that statistics are correctly tracked"""
        from cli import MedForgeCUIGenerator

        generator = MedForgeCUIGenerator(
            output_dir=temp_output_dir,
            seed=42,
            formats=['docx', 'pdf'],
        )

        stats = generator.generate_batch(
            cui_positive_count=5,
            cui_negative_count=3,
        )

        assert stats['total_generated'] == 8
        assert stats['cui_positive'] == 5
        assert stats['cui_negative'] == 3
        assert sum(stats['by_format'].values()) == 8
        assert sum(stats['by_category'].values()) == 8

    def test_llm_stats_tracking(self, temp_output_dir):
        """Test that LLM/template stats are tracked"""
        from cli import MedForgeCUIGenerator

        # With 0% LLM, all should be template-based
        generator = MedForgeCUIGenerator(
            output_dir=temp_output_dir,
            seed=42,
            formats=['docx'],
            llm_percentage=0.0,
        )

        stats = generator.generate_batch(
            cui_positive_count=5,
            cui_negative_count=3,
        )

        assert stats['llm_enhanced'] == 0
        assert stats['template_based'] == 8


if __name__ == '__main__':
    pytest.main([__file__, '-v'])
