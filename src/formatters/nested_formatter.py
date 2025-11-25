"""
Nested document formatter for emails with attachments
Creates complex scenarios: emails containing PHI documents
"""
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email.mime.application import MIMEApplication
from email import encoders
from datetime import datetime
import os
import random
import io
import zipfile
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.lib import colors
from docx import Document
from docx.shared import Pt, Inches


class NestedEmailFormatter:
    """Creates emails with document attachments (nested PHI scenarios)"""

    def __init__(self, output_dir='output'):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def create_email_with_lab_attachment(self, patient, provider, lab_pdf_path, filename):
        """
        Create email with lab result PDF attached (PHI Positive)
        This tests Purview's ability to detect PHI in nested documents
        """
        msg = MIMEMultipart('mixed')

        # Email headers
        msg['Subject'] = f"Lab Results - {patient['first_name']} {patient['last_name']}"
        msg['From'] = f"{provider['first_name']} {provider['last_name']} <{provider['email']}>"
        msg['To'] = f"{patient['first_name']} {patient['last_name']} <{patient['email']}>"
        msg['Date'] = datetime.now().strftime('%a, %d %b %Y %H:%M:%S %z')
        msg['Message-ID'] = f"<{random.randint(100000, 999999)}@healthsystem.org>"

        # Email body
        body_text = f"""
Dear {patient['first_name']} {patient['last_name']},

Please find attached your recent laboratory results from our office visit on {datetime.now().strftime('%m/%d/%Y')}.

Patient: {patient['first_name']} {patient['last_name']}
MRN: {patient['mrn']}
Date of Birth: {patient['dob'].strftime('%m/%d/%Y')}

The attached PDF contains your complete test results. Please review and contact our office if you have any questions.

Best regards,

{provider['first_name']} {provider['last_name']}, {provider['title']}
{provider['specialty']}
Phone: {provider['phone']}

---
CONFIDENTIAL: This email contains protected health information (PHI).
Unauthorized disclosure or forwarding is prohibited under HIPAA regulations.
"""

        body = MIMEText(body_text, 'plain')
        msg.attach(body)

        # Attach the PDF file
        if os.path.exists(lab_pdf_path):
            with open(lab_pdf_path, 'rb') as f:
                pdf_attachment = MIMEApplication(f.read(), _subtype='pdf')
                pdf_attachment.add_header(
                    'Content-Disposition',
                    'attachment',
                    filename=os.path.basename(lab_pdf_path)
                )
                msg.attach(pdf_attachment)

        # Save as EML
        filepath = os.path.join(self.output_dir, filename)
        with open(filepath, 'w') as f:
            f.write(msg.as_string())

        return filepath

    def create_email_with_multiple_attachments(self, patient, provider, attachment_paths, filename):
        """
        Create email with multiple document attachments (PHI Positive)
        Tests detection of PHI across multiple nested files
        """
        msg = MIMEMultipart('mixed')

        # Email headers
        msg['Subject'] = f"Medical Records - {patient['last_name']}, {patient['first_name']}"
        msg['From'] = "Medical Records <records@healthsystem.org>"
        msg['To'] = f"{provider['first_name']} {provider['last_name']} <{provider['email']}>"
        msg['Date'] = datetime.now().strftime('%a, %d %b %Y %H:%M:%S %z')
        msg['Message-ID'] = f"<{random.randint(100000, 999999)}@healthsystem.org>"

        # Email body
        body_text = f"""
Dr. {provider['last_name']},

As requested, please find attached medical records for:

Patient Name: {patient['first_name']} {patient['last_name']}
Date of Birth: {patient['dob'].strftime('%m/%d/%Y')}
MRN: {patient['mrn']}
Phone: {patient['phone']}
Address: {patient['address']['street']}, {patient['address']['city']}, {patient['address']['state']} {patient['address']['zip']}

Attached documents include:
{self._list_attachments(attachment_paths)}

Please confirm receipt of these records.

Medical Records Department
Phone: (555) 123-4567

---
CONFIDENTIAL MEDICAL RECORDS: This email and all attachments contain protected health information.
Handle in accordance with HIPAA privacy regulations.
"""

        body = MIMEText(body_text, 'plain')
        msg.attach(body)

        # Attach all files
        for attachment_path in attachment_paths:
            if os.path.exists(attachment_path):
                self._attach_file(msg, attachment_path)

        # Save as EML
        filepath = os.path.join(self.output_dir, filename)
        with open(filepath, 'w') as f:
            f.write(msg.as_string())

        return filepath

    def create_referral_email_with_notes(self, patient, referring_provider, specialist_provider,
                                        progress_note_path, filename):
        """
        Create referral email with progress note attached (PHI Positive)
        Common real-world scenario: provider sends patient records to specialist
        """
        msg = MIMEMultipart('mixed')

        # Email headers
        msg['Subject'] = f"Patient Referral: {patient['last_name']}, {patient['first_name']}"
        msg['From'] = f"{referring_provider['first_name']} {referring_provider['last_name']} <{referring_provider['email']}>"
        msg['To'] = f"{specialist_provider['first_name']} {specialist_provider['last_name']} <{specialist_provider['email']}>"
        msg['Date'] = datetime.now().strftime('%a, %d %b %Y %H:%M:%S %z')
        msg['Message-ID'] = f"<{random.randint(100000, 999999)}@healthsystem.org>"

        # Email body
        diagnoses_text = '\n'.join([f"- {d['name']} (ICD-10: {d['icd10']})" for d in patient['diagnoses']])
        meds_text = '\n'.join([f"- {med}" for med in patient['medications'][:5]])

        body_text = f"""
Dear Dr. {specialist_provider['last_name']},

I am referring the following patient to you for evaluation and management of {patient['diagnoses'][0]['name']}:

PATIENT INFORMATION:
Name: {patient['first_name']} {patient['last_name']}
DOB: {patient['dob'].strftime('%m/%d/%Y')} (Age: {patient['age']})
MRN: {patient['mrn']}
Contact: {patient['phone']} / {patient['email']}
Address: {patient['address']['street']}, {patient['address']['city']}, {patient['address']['state']} {patient['address']['zip']}

CURRENT DIAGNOSES:
{diagnoses_text}

CURRENT MEDICATIONS:
{meds_text}

ALLERGIES: {', '.join(patient['allergies'])}

REASON FOR REFERRAL:
Patient requires specialist evaluation for management of {patient['diagnoses'][0]['name']}.
Recent symptoms and lab findings suggest need for advanced treatment planning.

Attached is my most recent progress note with additional clinical details.

Please contact me if you need any additional information.

Thank you for seeing this patient.

Best regards,

{referring_provider['first_name']} {referring_provider['last_name']}, {referring_provider['title']}
{referring_provider['specialty']}
NPI: {referring_provider['npi']}
Phone: {referring_provider['phone']}
Fax: {referring_provider['fax']}
"""

        body = MIMEText(body_text, 'plain')
        msg.attach(body)

        # Attach progress note
        if os.path.exists(progress_note_path):
            self._attach_file(msg, progress_note_path)

        # Save as EML
        filepath = os.path.join(self.output_dir, filename)
        with open(filepath, 'w') as f:
            f.write(msg.as_string())

        return filepath

    def create_email_with_blank_form(self, facility, form_path, filename):
        """
        Create email with blank form template attached (PHI Negative)
        No patient data - just distributing forms
        """
        msg = MIMEMultipart('mixed')

        # Email headers
        msg['Subject'] = "Updated Patient Registration Forms"
        msg['From'] = f"Office Manager <manager@{facility['name'].lower().replace(' ', '')}.org>"
        msg['To'] = f"Front Desk Staff <frontdesk@{facility['name'].lower().replace(' ', '')}.org>"
        msg['Date'] = datetime.now().strftime('%a, %d %b %Y %H:%M:%S %z')
        msg['Message-ID'] = f"<{random.randint(100000, 999999)}@healthsystem.org>"

        # Email body
        body_text = f"""
Dear Front Desk Team,

Please find attached the updated patient registration forms effective January 1, 2025.

KEY UPDATES:
- Added emergency contact section
- Updated insurance information fields
- New HIPAA consent language
- Spanish language version included

Please print copies for the front desk and use these forms for all new patient registrations starting next month.

The old forms should be discarded after December 31, 2024.

If you have questions, please contact the Office Manager.

Thank you,

Office Management
{facility['name']}
{facility['address']['street']}
{facility['address']['city']}, {facility['address']['state']} {facility['address']['zip']}
Phone: {facility['phone']}
"""

        body = MIMEText(body_text, 'plain')
        msg.attach(body)

        # Attach blank form
        if os.path.exists(form_path):
            self._attach_file(msg, form_path)

        # Save as EML
        filepath = os.path.join(self.output_dir, filename)
        with open(filepath, 'w') as f:
            f.write(msg.as_string())

        return filepath

    def create_policy_email_with_pdf(self, facility, policy_pdf_path, filename):
        """
        Create policy distribution email with PDF attached (PHI Negative)
        Administrative content, no patient data
        """
        msg = MIMEMultipart('mixed')

        # Email headers
        msg['Subject'] = "New Clinical Documentation Policy - Action Required"
        msg['From'] = f"Compliance Department <compliance@{facility['name'].lower().replace(' ', '')}.org>"
        msg['To'] = f"All Clinical Staff <clinical@{facility['name'].lower().replace(' ', '')}.org>"
        msg['Date'] = datetime.now().strftime('%a, %d %b %Y %H:%M:%S %z')
        msg['Message-ID'] = f"<{random.randint(100000, 999999)}@healthsystem.org>"

        # Email body
        body_text = f"""
Dear Clinical Team,

Please review the attached Clinical Documentation Policy (CPG-2024-001) which becomes effective immediately.

SUMMARY OF CHANGES:
- Enhanced medication reconciliation requirements
- Updated electronic signature procedures
- New patient education documentation standards
- ICD-11 transition timeline

ACTION REQUIRED:
1. Read the attached policy document
2. Complete acknowledgment form by December 15, 2024
3. Attend one mandatory training session (dates in policy)

TRAINING SESSIONS:
- Session 1: January 10, 2025 at 2:00 PM
- Session 2: January 12, 2025 at 9:00 AM

Register for training through the staff portal.

Questions should be directed to the Compliance Department at ext. 2500.

Thank you for your cooperation.

Compliance Department
{facility['name']}
Phone: {facility['phone']}
"""

        body = MIMEText(body_text, 'plain')
        msg.attach(body)

        # Attach policy PDF
        if os.path.exists(policy_pdf_path):
            self._attach_file(msg, policy_pdf_path)

        # Save as EML
        filepath = os.path.join(self.output_dir, filename)
        with open(filepath, 'w') as f:
            f.write(msg.as_string())

        return filepath

    def _attach_file(self, msg, filepath):
        """Helper method to attach a file to an email message"""
        # Determine MIME type based on file extension
        ext = os.path.splitext(filepath)[1].lower()

        if ext == '.pdf':
            subtype = 'pdf'
            maintype = 'application'
        elif ext == '.docx':
            subtype = 'vnd.openxmlformats-officedocument.wordprocessingml.document'
            maintype = 'application'
        elif ext == '.xlsx':
            subtype = 'vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            maintype = 'application'
        elif ext == '.pptx':
            subtype = 'vnd.openxmlformats-officedocument.presentationml.presentation'
            maintype = 'application'
        else:
            subtype = 'octet-stream'
            maintype = 'application'

        with open(filepath, 'rb') as f:
            if maintype == 'application':
                attachment = MIMEApplication(f.read(), _subtype=subtype)
            else:
                attachment = MIMEBase(maintype, subtype)
                attachment.set_payload(f.read())
                encoders.encode_base64(attachment)

            attachment.add_header(
                'Content-Disposition',
                'attachment',
                filename=os.path.basename(filepath)
            )
            msg.attach(attachment)

    def _list_attachments(self, attachment_paths):
        """Helper to create a bullet list of attachment filenames"""
        return '\n'.join([f"- {os.path.basename(path)}" for path in attachment_paths])

    # ========================================
    # NEW: In-memory attachment generation
    # ========================================

    def _generate_phi_positive_pdf_in_memory(self, patient, provider, lab_data):
        """Generate PHI positive lab result PDF in memory"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        # Header
        story.append(Paragraph(f"<b>{provider['specialty']} - Laboratory Results</b>", styles['Title']))
        story.append(Spacer(1, 0.2*inch))

        # Patient info
        patient_data = [
            ['Patient Information', ''],
            ['Name:', f"{patient['first_name']} {patient['last_name']}"],
            ['MRN:', patient['mrn']],
            ['Date of Birth:', patient['dob'].strftime('%m/%d/%Y')],
            ['Collection Date:', datetime.now().strftime('%m/%d/%Y')],
        ]
        patient_table = Table(patient_data, colWidths=[2*inch, 4*inch])
        patient_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(patient_table)
        story.append(Spacer(1, 0.3*inch))

        # Lab results - handle both dict format (with 'results' key) and list format
        lab_results = lab_data.get('results', lab_data) if isinstance(lab_data, dict) else lab_data
        lab_table_data = [['Test', 'Result', 'Reference Range', 'Flag']]
        for test in lab_results:
            # Handle both 'name' and 'test' keys for test name
            test_name = test.get('name') or test.get('test', 'Unknown')
            lab_table_data.append([
                test_name,
                f"{test['value']} {test['unit']}",
                test['reference_range'],
                test.get('flag', '')
            ])

        lab_table = Table(lab_table_data, colWidths=[2.5*inch, 1.5*inch, 1.5*inch, 0.5*inch])
        lab_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightblue),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(lab_table)

        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()

    def _generate_phi_negative_pdf_in_memory(self, facility):
        """Generate PHI negative policy PDF in memory"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        # Header
        story.append(Paragraph(f"<b>{facility['name']}</b>", styles['Title']))
        story.append(Paragraph("<b>Clinical Documentation Policy</b>", styles['Heading1']))
        story.append(Spacer(1, 0.3*inch))

        # Policy content (no patient data)
        policy_text = """
        <b>POLICY NUMBER:</b> CPG-2024-001<br/>
        <b>EFFECTIVE DATE:</b> January 1, 2025<br/>
        <b>DEPARTMENT:</b> Clinical Operations<br/>
        <br/>
        <b>PURPOSE:</b><br/>
        To establish standardized procedures for clinical documentation across all departments.<br/>
        <br/>
        <b>SCOPE:</b><br/>
        This policy applies to all clinical staff including physicians, nurses, and allied health professionals.<br/>
        <br/>
        <b>POLICY STATEMENTS:</b><br/>
        1. All clinical encounters must be documented within 24 hours<br/>
        2. Electronic signatures are required for all clinical notes<br/>
        3. Medication reconciliation must be completed at each visit<br/>
        4. Patient education must be documented with standardized templates<br/>
        <br/>
        <b>COMPLIANCE:</b><br/>
        Failure to comply with this policy may result in disciplinary action.
        """
        story.append(Paragraph(policy_text, styles['Normal']))

        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()

    def _generate_phi_positive_docx_in_memory(self, patient, provider):
        """Generate PHI positive progress note DOCX in memory"""
        buffer = io.BytesIO()
        doc = Document()

        # Header
        doc.add_heading(f'Progress Note - {patient["first_name"]} {patient["last_name"]}', 0)

        # Patient info table
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'

        cells = table.rows[0].cells
        cells[0].text = 'Patient Name:'
        cells[1].text = f"{patient['first_name']} {patient['last_name']}"

        cells = table.rows[1].cells
        cells[0].text = 'MRN:'
        cells[1].text = patient['mrn']

        cells = table.rows[2].cells
        cells[0].text = 'Date of Birth:'
        cells[1].text = patient['dob'].strftime('%m/%d/%Y')

        cells = table.rows[3].cells
        cells[0].text = 'Visit Date:'
        cells[1].text = datetime.now().strftime('%m/%d/%Y')

        cells = table.rows[4].cells
        cells[0].text = 'Provider:'
        cells[1].text = f"{provider['first_name']} {provider['last_name']}, {provider['title']}"

        doc.add_paragraph()

        # SOAP note
        doc.add_heading('Subjective:', level=2)
        # Get first diagnosis safely
        first_diagnosis = "chronic condition"
        if 'diagnoses' in patient and len(patient['diagnoses']) > 0:
            if isinstance(patient['diagnoses'][0], dict):
                first_diagnosis = patient['diagnoses'][0].get('name', 'chronic condition')
            else:
                first_diagnosis = str(patient['diagnoses'][0])

        doc.add_paragraph(f"Patient presents for follow-up of {first_diagnosis}. "
                         f"Reports adherence to current medication regimen.")

        doc.add_heading('Objective:', level=2)
        # Handle vital signs
        if 'vital_signs' in patient and isinstance(patient['vital_signs'], dict):
            vitals = patient['vital_signs']
            vitals_text = f"BP: {vitals.get('blood_pressure', '120/80')}, HR: {vitals.get('heart_rate', '72')}, "
            vitals_text += f"Temp: {vitals.get('temperature', '98.6')}°F"
        else:
            vitals_text = "BP: 120/80, HR: 72, Temp: 98.6°F"
        doc.add_paragraph(vitals_text)

        doc.add_heading('Assessment:', level=2)
        doc.add_paragraph(f"1. {first_diagnosis} - stable")

        doc.add_heading('Plan:', level=2)
        # Get first medication safely
        first_med = "current medications"
        if 'medications' in patient and len(patient['medications']) > 0:
            first_med = patient['medications'][0]
        doc.add_paragraph(f"Continue {first_med}")
        doc.add_paragraph("Follow-up in 3 months or sooner if symptoms worsen.")

        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _generate_phi_negative_docx_in_memory(self, facility):
        """Generate PHI negative blank form DOCX in memory"""
        buffer = io.BytesIO()
        doc = Document()

        doc.add_heading(f'{facility["name"]} - Patient Registration Form', 0)
        doc.add_paragraph('Please complete all fields below:')
        doc.add_paragraph()

        # Blank form fields (no actual patient data)
        doc.add_heading('Patient Information', level=2)
        doc.add_paragraph('Last Name: _______________________________')
        doc.add_paragraph('First Name: _______________________________')
        doc.add_paragraph('Date of Birth: ____/____/________')
        doc.add_paragraph('Phone Number: (_____) _____-_______')
        doc.add_paragraph('Email: _______________________________')
        doc.add_paragraph()

        doc.add_heading('Emergency Contact', level=2)
        doc.add_paragraph('Name: _______________________________')
        doc.add_paragraph('Relationship: _______________________________')
        doc.add_paragraph('Phone: (_____) _____-_______')
        doc.add_paragraph()

        doc.add_heading('Insurance Information', level=2)
        doc.add_paragraph('Insurance Company: _______________________________')
        doc.add_paragraph('Policy Number: _______________________________')
        doc.add_paragraph('Group Number: _______________________________')

        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _create_zip_with_phi_positive_docs(self, patient, provider, lab_data):
        """Create ZIP file with 2-3 PHI positive documents in memory"""
        buffer = io.BytesIO()

        with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
            # Add lab result PDF
            lab_pdf = self._generate_phi_positive_pdf_in_memory(patient, provider, lab_data)
            zipf.writestr(f"LabResults_{patient['mrn']}.pdf", lab_pdf)

            # Add progress note DOCX
            progress_note = self._generate_phi_positive_docx_in_memory(patient, provider)
            zipf.writestr(f"ProgressNote_{patient['mrn']}.docx", progress_note)

            # Randomly add a third document (50% chance)
            if random.random() < 0.5:
                # Add another lab result with different data
                zipf.writestr(f"PreviousLab_{patient['mrn']}.pdf", lab_pdf)

        buffer.seek(0)
        return buffer.getvalue()

    def _create_zip_with_phi_negative_docs(self, facility):
        """Create ZIP file with 2-3 PHI negative documents in memory"""
        buffer = io.BytesIO()

        with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
            # Add policy PDF
            policy_pdf = self._generate_phi_negative_pdf_in_memory(facility)
            zipf.writestr("ClinicalDocumentationPolicy.pdf", policy_pdf)

            # Add blank form DOCX
            blank_form = self._generate_phi_negative_docx_in_memory(facility)
            zipf.writestr("PatientRegistrationForm.docx", blank_form)

            # Randomly add a third document (50% chance)
            if random.random() < 0.5:
                zipf.writestr("HIPAAConsentForm.docx", blank_form)

        buffer.seek(0)
        return buffer.getvalue()

    def create_phi_positive_email_with_attachment(self, patient, provider, facility, lab_data, filename):
        """
        Create PHI POSITIVE email with embedded attachment (PDF or ZIP)
        20% chance of ZIP, 80% chance of single PDF/DOCX
        """
        msg = MIMEMultipart('mixed')

        # Email headers with PHI
        msg['Subject'] = f"Lab Results - {patient['first_name']} {patient['last_name']}"
        msg['From'] = f"{provider['first_name']} {provider['last_name']} <{provider['email']}>"
        msg['To'] = f"{patient['first_name']} {patient['last_name']} <{patient['email']}>"
        msg['Date'] = datetime.now().strftime('%a, %d %b %Y %H:%M:%S %z')
        msg['Message-ID'] = f"<{random.randint(100000, 999999)}@healthsystem.org>"

        # Email body with PHI
        body_text = f"""
Dear {patient['first_name']} {patient['last_name']},

Please find attached your recent laboratory results.

Patient: {patient['first_name']} {patient['last_name']}
MRN: {patient['mrn']}
Date of Birth: {patient['dob'].strftime('%m/%d/%Y')}

Please review and contact our office if you have any questions.

Best regards,

{provider['first_name']} {provider['last_name']}, {provider['title']}
{provider['specialty']}
Phone: {provider['phone']}

---
CONFIDENTIAL: This email contains protected health information (PHI).
"""

        body = MIMEText(body_text, 'plain')
        msg.attach(body)

        # Decide on attachment type (20% ZIP, 80% single doc)
        use_zip = random.random() < 0.2

        if use_zip:
            # Attach ZIP with multiple PHI documents
            zip_data = self._create_zip_with_phi_positive_docs(patient, provider, lab_data)
            attachment = MIMEApplication(zip_data, _subtype='zip')
            attachment.add_header('Content-Disposition', 'attachment',
                                filename=f"MedicalRecords_{patient['mrn']}.zip")
            msg.attach(attachment)
        else:
            # Attach single PDF or DOCX (50/50 split)
            if random.random() < 0.5:
                # PDF lab result
                pdf_data = self._generate_phi_positive_pdf_in_memory(patient, provider, lab_data)
                attachment = MIMEApplication(pdf_data, _subtype='pdf')
                attachment.add_header('Content-Disposition', 'attachment',
                                    filename=f"LabResults_{patient['mrn']}.pdf")
            else:
                # DOCX progress note
                docx_data = self._generate_phi_positive_docx_in_memory(patient, provider)
                attachment = MIMEApplication(docx_data,
                    _subtype='vnd.openxmlformats-officedocument.wordprocessingml.document')
                attachment.add_header('Content-Disposition', 'attachment',
                                    filename=f"ProgressNote_{patient['mrn']}.docx")
            msg.attach(attachment)

        # Save as EML
        filepath = os.path.join(self.output_dir, filename)
        with open(filepath, 'w') as f:
            f.write(msg.as_string())

        return filepath

    def create_phi_negative_email_with_attachment(self, facility, filename):
        """
        Create PHI NEGATIVE email with embedded attachment (PDF or ZIP)
        20% chance of ZIP, 80% chance of single PDF/DOCX
        """
        msg = MIMEMultipart('mixed')

        # Email headers with NO patient data
        msg['Subject'] = "Updated Clinical Documentation Policy"
        msg['From'] = f"Compliance <compliance@{facility['name'].lower().replace(' ', '')}.org>"
        msg['To'] = f"All Staff <staff@{facility['name'].lower().replace(' ', '')}.org>"
        msg['Date'] = datetime.now().strftime('%a, %d %b %Y %H:%M:%S %z')
        msg['Message-ID'] = f"<{random.randint(100000, 999999)}@healthsystem.org>"

        # Email body with NO patient data
        body_text = f"""
Dear Team,

Please find attached the updated clinical documentation policy and forms.

This policy is effective January 1, 2025. All staff must review and acknowledge by December 31, 2024.

Training sessions will be scheduled in the coming weeks.

Thank you,

Compliance Department
{facility['name']}
Phone: {facility['phone']}
"""

        body = MIMEText(body_text, 'plain')
        msg.attach(body)

        # Decide on attachment type (20% ZIP, 80% single doc)
        use_zip = random.random() < 0.2

        if use_zip:
            # Attach ZIP with multiple PHI NEGATIVE documents
            zip_data = self._create_zip_with_phi_negative_docs(facility)
            attachment = MIMEApplication(zip_data, _subtype='zip')
            attachment.add_header('Content-Disposition', 'attachment',
                                filename="PolicyDocuments.zip")
            msg.attach(attachment)
        else:
            # Attach single PDF or DOCX (50/50 split)
            if random.random() < 0.5:
                # PDF policy
                pdf_data = self._generate_phi_negative_pdf_in_memory(facility)
                attachment = MIMEApplication(pdf_data, _subtype='pdf')
                attachment.add_header('Content-Disposition', 'attachment',
                                    filename="ClinicalDocumentationPolicy.pdf")
            else:
                # DOCX blank form
                docx_data = self._generate_phi_negative_docx_in_memory(facility)
                attachment = MIMEApplication(docx_data,
                    _subtype='vnd.openxmlformats-officedocument.wordprocessingml.document')
                attachment.add_header('Content-Disposition', 'attachment',
                                    filename="PatientRegistrationForm.docx")
            msg.attach(attachment)

        # Save as EML
        filepath = os.path.join(self.output_dir, filename)
        with open(filepath, 'w') as f:
            f.write(msg.as_string())

        return filepath
