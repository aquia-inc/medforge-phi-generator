"""
PDF Form Field Populator

Populates fillable PDF forms with synthetic data using Faker.
Works with customer-provided CMS template forms.
"""
import logging
import pikepdf
from pikepdf import Pdf
from pikepdf.form import Form
from faker import Faker
import random
import os
from datetime import datetime
from typing import Dict, Any, Optional

logger = logging.getLogger(__name__)


class PDFFormPopulator:
    """Populates fillable PDF forms with synthetic data."""

    def __init__(self, seed: Optional[int] = None):
        """Initialize with optional random seed."""
        self.fake = Faker('en_US')
        if seed:
            Faker.seed(seed)
            random.seed(seed)

    def populate_acroform(self, template_path: str, output_path: str,
                          field_data: Dict[str, Any]) -> str:
        """
        Populate a PDF AcroForm with synthetic data using pikepdf.form.Form.

        Uses the modern pikepdf Form API which handles NeedAppearances internally,
        ensuring SharePoint/Purview render and index the filled content correctly.

        Args:
            template_path: Path to AcroForm PDF template
            output_path: Path to save populated PDF
            field_data: Dictionary mapping AcroForm field names to values.
                        bool True → checked, bool False → unchecked.

        Returns:
            Path to created file
        """
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        try:
            with Pdf.open(template_path) as pdf:
                form = Form(pdf)
                for field_name, value in field_data.items():
                    if field_name in form:
                        if isinstance(value, bool):
                            form[field_name].checked = value
                        else:
                            form[field_name].value = str(value) if value else ''
                pdf.save(output_path)
        except Exception as e:
            raise RuntimeError(
                f"pikepdf AcroForm population failed for {template_path}: {e}"
            ) from e

        return output_path

    def populate_form(self, template_path: str, output_path: str, field_data: Dict[str, Any],
                      field_positions: Optional[Dict[str, tuple]] = None) -> str:
        """
        Populate a flat PDF (no AcroForm) with a reportlab text overlay.

        For AcroForm PDFs (fillable forms with named fields), use populate_acroform()
        instead — it uses pikepdf with NeedAppearances=True which SharePoint/Purview
        render and index correctly.

        Args:
            template_path: Path to blank PDF template
            output_path: Path to save populated PDF
            field_data: Dictionary mapping field names to values
            field_positions: Optional dict mapping field names to (x, y) coordinates.
                           If None, uses default Reasonable Accommodation positions.

        Returns:
            Path to created file
        """
        # Create output directory
        os.makedirs(os.path.dirname(output_path), exist_ok=True)

        # Use reportlab to overlay text on template PDF (only way that renders everywhere)
        primary_error = None
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            from io import BytesIO
            from PyPDF2 import PdfReader, PdfWriter

            # Create overlay with text
            packet = BytesIO()
            can = canvas.Canvas(packet, pagesize=letter)

            # Use provided positions or default Reasonable Accommodation form positions
            if field_positions is None:
                field_positions = {
                    'Name': (145, 620),
                    'Component': (145, 596),
                    'Telephone number': (180, 574),
                    'Location': (200, 548),
                    'Grade': (150, 525),
                    'Date of Birth': (147, 501),
                    'Manager': (130, 476),
                    'Discription': (88, 380),  # Large text area
                }

            # Draw text on PDF
            can.setFont("Helvetica", 10)
            for field_name, value in field_data.items():
                if field_name in field_positions and value and value not in [True, False]:
                    x, y = field_positions[field_name]
                    can.drawString(x, y, str(value)[:80])

            can.save()
            packet.seek(0)

            # Overlay on template
            overlay = PdfReader(packet)
            template = PdfReader(template_path)
            output = PdfWriter()

            # Merge overlay with template
            page = template.pages[0]
            page.merge_page(overlay.pages[0])
            output.add_page(page)

            # Add remaining pages if any
            for i in range(1, len(template.pages)):
                output.add_page(template.pages[i])

            # Write
            with open(output_path, 'wb') as f:
                output.write(f)

        except ImportError as e:
            logger.warning("reportlab/PyPDF2 not installed, using pikepdf fallback: %s", e)
            primary_error = e
        except Exception as e:
            logger.warning("reportlab overlay failed for %s: %s", template_path, e)
            primary_error = e

        if primary_error is None:
            return output_path

        # Fallback: reached only when reportlab overlay fails (import error or exception).
        # For AcroForm PDFs, call populate_acroform() directly instead of this method.
        try:
            pdf = pikepdf.open(template_path)

            if '/AcroForm' in pdf.Root and '/Fields' in pdf.Root.AcroForm:
                for field in pdf.Root.AcroForm.Fields:
                    field_name = str(field.T) if '/T' in field else None

                    if field_name and field_name in field_data:
                        value = field_data[field_name]

                        if value is True:
                            field['/V'] = pikepdf.Name('/On')
                        elif value is False:
                            field['/V'] = pikepdf.Name('/Off')
                        else:
                            field['/V'] = str(value) if value else ''

                        if '/AP' in field:
                            del field['/AP']

                pdf.Root.AcroForm['/NeedAppearances'] = True

            pdf.save(output_path)
            pdf.close()
            logger.warning(
                "Used pikepdf fallback for %s -- form field data may not be "
                "extractable by Purview/SharePoint text indexers", output_path
            )

        except Exception as e2:
            logger.error(
                "Both reportlab and pikepdf failed for %s: %s / %s",
                template_path, primary_error, e2
            )
            raise RuntimeError(
                f"Cannot populate PDF form {template_path}: "
                f"reportlab error: {primary_error}, pikepdf error: {e2}"
            ) from e2

        return output_path

    def generate_medical_inquiry_data(self) -> Dict[str, Any]:
        """Generate data for Medical Inquiry Form (PHI)."""

        # Generate employee/patient info
        first_name = self.fake.first_name()
        last_name = self.fake.last_name()
        employee_name = f"{first_name} {last_name}"

        # Medical impairment options
        impairments = [
            "Severe latex allergy with contact dermatitis and respiratory symptoms",
            "Chronic lower back pain with limited mobility and sitting tolerance",
            "Type 1 Diabetes requiring insulin management and dietary modifications",
            "Severe asthma requiring inhaler use and environmental controls",
            "Rheumatoid arthritis affecting hand function and fine motor tasks",
            "Hearing loss requiring hearing aids and communication accommodations",
            "Visual impairment requiring screen reader and magnification software",
            "Chronic migraine disorder triggered by fluorescent lighting and stress",
        ]

        impairment = random.choice(impairments)

        # Duration options
        durations = ["permanent", "6 months", "1 year", "2 years", "indefinite"]
        duration = random.choice(durations)

        # Suggestions for accommodations
        accommodation_suggestions = [
            "Modified work schedule, ergonomic workstation, alternative lighting",
            "Flexible break schedule, accessible workspace location",
            "Remote work option 2-3 days per week, modified hours",
            "Assistive technology, adjusted performance standards",
            "Environmental modifications, alternative duty assignments",
        ]

        # Provider info
        provider_name = f"Dr. {self.fake.last_name()}, MD"

        # Major life activities - randomly select 2-4
        activities = {
            'Caring For Self': False,
            'Walking': False,
            'Hearing': False,
            'Lifting': False,
            'Interacting With Others': False,
            'Standing': False,
            'Seeing': False,
            'Sleeping': False,
            'Performing Manual Tasks': False,
            'Reaching': False,
            'Speaking': False,
            'Concentrating': False,
            'Breathing': False,
            'Thinking': False,
            'Learning': False,
            'Reproduction': False,
            'Working': False,
            'Toileting': False,
            'Sitting': False,
        }

        # Force 2-4 activities to be selected
        activities_list = list(activities.keys())
        for _ in range(random.randint(2, 4)):
            activities[random.choice(activities_list)] = True

        form_data = {
            'Employee Name Click here to enter text': employee_name,
            'Does the employee have a physical or mental impairment': 'Yes',
            'What is the impairmentdiagnosis Click here to enter text': impairment,
            'What is the expected duration of the impairment x months x years or permanent Click here to enter text': duration,
            'Does the impairment affect a major life activity': 'Yes_2',
            'Please describe how the employees limitations interfere with their ability to perform the job functions Click here to enter text':
                f"The employee's {impairment.split()[0].lower()} condition significantly impacts their ability to perform essential job functions without accommodation.",
            'Do you have any suggestions regarding possible accommodations to improve job performance  If so what are they Click here to enter text':
                random.choice(accommodation_suggestions),
            'If you have any additional comments please include them below Click here to enter text':
                "Employee is motivated and capable of performing job duties with reasonable accommodations in place.",
            'Print Name': provider_name,
            'Date': datetime.now().strftime('%m/%d/%Y'),
        }

        # Add activity checkboxes
        form_data.update(activities)

        return form_data

    def generate_eft_authorization_data(self) -> Dict[str, Any]:
        """Generate data for EFT Authorization Form (CUI-Finance)."""

        company_name = self.fake.company()
        contact_name = self.fake.name()

        # Generate routing number (9 digits, must be valid checksum)
        routing_number = f"{random.randint(100000000, 999999999)}"

        # Generate account number (8-12 digits)
        account_number = f"{random.randint(10000000, 999999999999)}"

        # Generate TIN/EIN (9 digits)
        tin = f"{random.randint(100000000, 999999999)}"

        # Generate UEI (12 character alphanumeric) - some vendors have this
        uei = ''.join(random.choices('ABCDEFGHJKLMNPQRSTUVWXYZ0123456789', k=12)) if random.random() < 0.3 else ''

        # Generate CAGE code (5 character alphanumeric) - procurement vendors
        cage = ''.join(random.choices('0123456789ABCDEFGHJKLMNPQRSTUVWXYZ', k=5)) if random.random() < 0.2 else ''

        form_data = {
            # Part 1: Account Holder Information
            'txtPayee': company_name,
            'txtDBA': '' if random.random() < 0.7 else self.fake.company_suffix(),
            'txtAHStreet': self.fake.street_address(),
            'txtAHCity': self.fake.city(),
            'txtAHState': self.fake.state_abbr(),
            'txtAHZip': self.fake.zipcode(),
            'txtTIN': tin,
            'txtTINType': random.choice(['SSN Individual', 'EIN Organization']),
            'txtUEI': uei,
            'txtCAGE': cage,
            'txtContactName': contact_name,
            'txtContactTelephone': self.fake.phone_number(),

            # Part 2: Financial Institution Information
            'txtBankName': random.choice([
                'Bank of America', 'Wells Fargo', 'Chase Bank', 'Citibank',
                'US Bank', 'PNC Bank', 'Capital One', 'TD Bank',
                'Truist Bank', 'Fifth Third Bank', 'Citizens Bank'
            ]),
            'txtRoutingNum': routing_number,
            'txtDepositNum': account_number,
            'txtTypeofAccount': random.choice(['Checking Account', 'Savings Account']),

            # Part 3: CMS Administrative Fields (use exact dropdown values from PDF)
            'Vendor Type': random.choice(['Customer', 'Supplier', 'Both - Cus. & Sup.']),
            'CMS Employee': random.choice(['Yes', 'No']),
            'SES Employee': random.choice(['Yes', 'No']),
            'Federal Vendor': random.choice(['Yes', 'No']),
            '1099': random.choice(['Yes', 'No']),
            'Trading Partner': random.choice(['Yes', 'No']),

            # Signature
            'txtSignature': contact_name,
        }

        return form_data

    # --- Critical Infrastructure generators ---

    # --- Shared helpers ---

    FISMA_SYSTEMS = [
        'CMS Cloud Services', 'Medicare Fee-for-Service', 'Healthcare.gov',
        'Marketplace Platform', 'Quality Payment Program', 'HCFAC System',
        'Enterprise Identity Management', 'CMS Data Exchange',
    ]

    SYSTEM_NAMES = [
        'CloudVault', 'SecureEdge', 'DataShield', 'NetGuard', 'CipherNet',
        'TrustCore', 'SafeLink', 'VaultStream', 'InfoSentry', 'CryptoGrid',
    ]

    CMS_OFFICES = ['OIT', 'CCIIO', 'CM', 'CMCS', 'OFM', 'OAGM', 'OA']

    def _foia_staff_names(self) -> tuple:
        """Return (director_name, liaison_name) for FOIA response letters."""
        return self.fake.name(), self.fake.name()

    def _rfc_memo_base_data(self, contract_placeholder: str) -> Dict[str, Any]:
        """Shared fill data for RFC Memo and AGX RFC Memo."""
        mod_type = random.choice(['Administrative', 'No-Cost Extension', 'Option Exercise',
                                   'Incremental Funding', 'Change Order'])
        return {
            'XX/XX/XXXX': datetime.now().strftime('%m/%d/%Y'),
            contract_placeholder:
                f"{self._contract_number()} / {self._task_order_number()}",
            'TBD': f"APP-{random.randint(2024, 2026)}-{random.randint(100, 999)}",
            '_underline_fills': [
                f"__X__ {mod_type}" if mod_type in ['Administrative', 'No-Cost Extension', 'Option Exercise'] else f"_____ {mod_type}",
                f"__X__ {mod_type}" if mod_type in ['Incremental Funding', 'Change Order'] else f"_____ {mod_type}",
                '',  # Negotiated Change line
                self.fake.name(),  # Signature
            ],
        }

    def generate_kmp_data(self) -> Dict[str, Any]:
        """Generate data for Key Management Plan (CUI-Critical Infrastructure)."""
        system = random.choice(self.SYSTEM_NAMES)
        system_lower = system.lower().replace(' ', '-')
        return {
            'Mock System': system,
            'Mock-system': f"{system_lower}",
            'MockSystem': system.replace(' ', ''),
            'Mock system': system,
        }

    def generate_rules_of_behavior_data(self) -> Dict[str, Any]:
        """Generate data for Rules of Behavior (CUI-Critical Infrastructure).

        The negative template has [MAC NAME] placeholders; positive is pre-filled.
        """
        MAC_NAMES = [
            'CCSQ MAC', 'FISS MAC', 'HIGLAS MAC', 'MCS MAC', 'VMS MAC',
            'BCRC MAC', 'EDPS MAC', 'HETS MAC', 'MFFS MAC', 'SPS MAC',
        ]
        mac = random.choice(MAC_NAMES)
        return {
            '[MAC NAME]': mac,
            'MAC NAME': mac,
        }

    def generate_hhs_rbd_data(self) -> Dict[str, Any]:
        """Generate data for HHS Rules of Behavior Deviation (CUI-Critical Infrastructure).

        Fills 57 PDF form fields: names, system info, risk justifications, approvals.
        """
        employee = self.fake.name()
        sso = self.fake.name()
        so = self.fake.name()
        opdiv_dir = self.fake.name()
        ciso = self.fake.name()
        ao = self.fake.name()
        cfo = self.fake.name()

        system_name = random.choice(self.FISMA_SYSTEMS)

        WEAKNESSES = [
            'Legacy system requires extended password expiration beyond 90-day policy',
            'Third-party application does not support multi-factor authentication',
            'Contractor VPN access requires broader network permissions than standard policy',
            'Cloud service provider audit logging does not meet NIST 800-53 AU-3 requirements',
            'System uses deprecated TLS 1.1 pending vendor migration to TLS 1.3',
        ]
        MITIGATIONS = [
            'Compensating controls include enhanced monitoring and quarterly access reviews',
            'Additional logging and alerting configured at network boundary',
            'Risk accepted with condition of quarterly review and 12-month remediation plan',
            'Dedicated security monitoring with automated anomaly detection in place',
            'Network segmentation isolates affected system from production environment',
        ]
        JUSTIFICATIONS = [
            'System upgrade scheduled for next fiscal year; interim deviation required for operations',
            'Vendor roadmap confirms compliance capability in next major release',
            'Mission-critical operations depend on this system; no alternative currently available',
            'Cost of immediate remediation exceeds approved budget; phased approach required',
        ]

        return {
            'S1: Name': employee,
            'S1: Email/Phone': f"{employee.split()[0].lower()}.{employee.split()[-1].lower()}@cms.hhs.gov",
            'S1: OpDiv Name': 'CMS',
            's1: System Name': system_name,
            's1: UUID': f"CMS-{random.randint(10000, 99999)}",
            'Request Duration': random.choice(['1 year', '2 years', '3 years']),
            'Program Area': random.choice(['CISO', 'OIT', 'CMCS', 'CCIIO', 'CM']),
            'System Security Officer Name': sso,
            'System Owner Name': so,
            's1: System Overview': f"{system_name} provides critical services for CMS operations and requires a deviation from standard security policy.",
            's2: Policy Directive': 'HHS Approved Physical Access and Logical Access Security Policy',
            's2: State the policy': f"The standard policy requires compliance with NIST 800-53 controls as implemented by {system_name}.",
            's3: Weakness Description': random.choice(WEAKNESSES),
            's3: Weakness Identifier': f"POA&M-{random.randint(2024, 2026)}-{random.randint(100, 999)}",
            's3: Risk Mitigation Controls': random.choice(MITIGATIONS),
            's3:Operational Justification': random.choice(JUSTIFICATIONS),
            's3: Plan for Compliance': 'Remediation plan targets full compliance within the approved deviation period.',
            's4: PII impact details': 'No direct impact to PII expected under this deviation.',
            's5: System Security Officer Name': sso,
            's5: SO Name': so,
            's6: OpDiv Director Name': opdiv_dir,
            's6: OpDiv Chief Financial Officer Name': cfo,
            's6: OpDiv Senior Official Name': opdiv_dir,
            's6: CISO Name': ciso,
            's6: AO Name': ao,
            'Security Request': True,
            'New Request': True,
        }

    def generate_incident_response_data(self) -> Dict[str, Any]:
        """Generate data for Incident Response Report (CUI-Critical Infrastructure).

        Fills tables with synthetic incident data: contacts, incident details, timeline.
        """
        reporter_first = self.fake.first_name()
        reporter_last = self.fake.last_name()
        reporter_email = f"{reporter_first}.{reporter_last}@cms.hhs.gov".lower()
        mgr_first = self.fake.first_name()
        mgr_last = self.fake.last_name()
        mgr_email = f"{mgr_first}.{mgr_last}@cms.hhs.gov".lower()
        fisma = random.choice(self.FISMA_SYSTEMS)
        CMS_OFFICES = self.CMS_OFFICES
        incident_date = self.fake.date_between(start_date='-60d', end_date='today')
        phone1 = self.fake.numerify('443-555-####')
        phone2 = self.fake.numerify('443-555-####')
        phone3 = self.fake.numerify('443-555-####')

        actions = [
            'Isolated affected systems and blocked malicious IP addresses',
            'Reviewed system logs and identified unauthorized access attempts',
            'Disabled compromised user accounts and reset credentials',
            'Deployed updated endpoint protection signatures',
            'Notified CISO and initiated forensic analysis',
        ]

        return {
            '_table_data': [
                # Table 0: Contact info (rows 2 and 4 have data)
                {
                    'table_index': 0,
                    'start_row': 2,
                    'rows': [
                        [reporter_first, reporter_last, 'Contractor', 'Contractor',
                         reporter_email, reporter_email],
                    ],
                },
                {
                    'table_index': 0,
                    'start_row': 4,
                    'rows': [
                        [phone1, phone2, phone3,
                         random.choice(self.CMS_OFFICES), random.choice(self.CMS_OFFICES),
                         f"E{random.randint(100, 999)}"],
                    ],
                },
                # Table 1: Impact counts
                {
                    'table_index': 1,
                    'start_row': 1,
                    'rows': [
                        [None, random.choice(['Yes', 'No'])],
                        [None, str(random.randint(1, 500))],
                        [None, str(random.randint(0, 100))],
                    ],
                },
                # Table 2: Timeline
                {
                    'table_index': 2,
                    'start_row': 1,
                    'rows': [
                        [None, incident_date.strftime('%b %d, %Y')],
                        [random.choice(actions), random.choice(actions)],
                        [None, incident_date.strftime('%b %d, %Y')],
                        [random.choice(actions), random.choice(actions)],
                        [None, incident_date.strftime('%b %d, %Y')],
                        ['Containment actions initiated per incident response plan.',
                         'Containment actions initiated per incident response plan.'],
                        [None, incident_date.strftime('%b %d, %Y')],
                        ['Recovery procedures in progress; monitoring for recurrence.',
                         'Recovery procedures in progress; monitoring for recurrence.'],
                    ],
                },
                # Table 3: System info
                {
                    'table_index': 3,
                    'start_row': 1,
                    'rows': [
                        [None, fisma, fisma, fisma, fisma, fisma, fisma],
                    ],
                },
                {
                    'table_index': 3,
                    'start_row': 4,
                    'rows': [
                        [None, mgr_first, mgr_last, 'Manager',
                         mgr_email, phone2, 'Y'],
                        [None, reporter_first, reporter_last, 'Analyst',
                         reporter_email, phone1, 'Y'],
                    ],
                },
            ],
        }

    # --- Financial generators ---

    def generate_afr_additional_info_data(self) -> Dict[str, Any]:
        """Generate data for Additional Info AFR (CUI-Financial)."""
        SECURITY_TOOLS = [
            'CrowdStrike Falcon', 'Palo Alto Cortex', 'Splunk Enterprise Security',
            'Tenable.io', 'Qualys VMDR', 'Carbon Black Cloud', 'SentinelOne Singularity',
        ]
        VENDORS = ['Accenture Federal', 'Deloitte', 'Booz Allen Hamilton',
                   'GDIT', 'Leidos', 'CGI Federal', 'Perspecta']
        tool = random.choice(SECURITY_TOOLS)
        vendor = random.choice(VENDORS)
        old_vendors = random.sample(['Symantec', 'McAfee', 'Trend Micro', 'FireEye', 'Forcepoint'], 2)
        return {
            'Mock Security Tool': tool,
            'Mock Security Tool from PRIORITY VENDOR': f"{tool} from {vendor}",
            'PRIORITY VENDOR': vendor,
            'Special Computers': random.choice(['CMS workstations', 'HHS endpoints', 'agency laptops',
                                                  'cloud instances', 'server infrastructure']),
            'OTHERVENDOR': random.choice(VENDORS),
            'OldVendor': old_vendors[0],
            'OldProduct': f"{old_vendors[1]} Endpoint Protection",
        }

    def generate_dibo_afr_data(self) -> Dict[str, Any]:
        """Generate data for DIBO AFR (CUI-Financial)."""
        CLOUD_PRODUCTS = [
            'CloudSecure', 'GovCloud Pro', 'FedConnect', 'DataVault Enterprise',
            'SecureHost', 'CipherStack', 'NetShield Pro', 'InfoGuard Cloud',
        ]
        product = random.choice(CLOUD_PRODUCTS)
        contract_num = f"HHSM-500-{random.randint(2030, 2036)}-{random.randint(10000, 99999)}I"
        task_order = self._contract_number()
        base_budget = random.randint(300000, 800000)
        base_cost = random.randint(1000000, 3000000)
        return {
            "HHSM-500-2034-00016I_75FCMC34R0002 BestCloud (BestCloud)":
                f"{contract_num}_{task_order} {product} ({product})",
            'BestCloud': product,
            "$428,293": self.format_currency(base_budget),
            "$1,868,293": self.format_currency(base_cost),
            "$1,440,000": self.format_currency(base_cost - base_budget),
            "$1,954,129": self.format_currency(base_cost * 1.046),
            "$1,525,836": self.format_currency(base_cost * 1.046 - base_budget),
            "$2,044,024": self.format_currency(base_cost * 1.094),
            "$1,615,731": self.format_currency(base_cost * 1.094 - base_budget),
        }

    def generate_supplemental_afr_data(self) -> Dict[str, Any]:
        """Generate data for Supplemental AFR (CUI-Financial)."""
        PRODUCTS = [
            'Zero Trust Gateway', 'Cloud Access Broker', 'Privileged Access Manager',
            'Threat Intelligence Platform', 'Container Security Suite',
            'API Gateway Manager', 'Identity Governance Platform',
        ]
        product = random.choice(PRODUCTS)
        return {
            'SuperSecure Support': f"{product} Support",
            'Internal Products and Service Now': f"{product} and ServiceNow",
            'Internal Products': product,
            'SSS solution': f"{self.fake.company()} solution",
            'SecurityKeys': random.choice(['YubiKeys', 'PIV cards', 'FIDO2 tokens', 'smart cards']),
        }

    def generate_oit_fo_aif_data(self) -> Dict[str, Any]:
        """Generate data for OIT FO Additional Info Form (CUI-Financial).

        Fills the blank questionnaire with realistic budget justification answers.
        """
        PRODUCTS = [
            'Cloud Migration Accelerator', 'AI-Powered Claims Processing',
            'Zero Trust Architecture Implementation', 'Enterprise Data Lake',
            'Automated Compliance Monitoring', 'Digital Identity Platform',
        ]
        product = random.choice(PRODUCTS)
        cost = self.generate_currency_amount(500000, 10000000)
        return {
            'Additional Funds Request - Additional Information': f"Additional Funds Request - {product}",
            # These appear after the question text on the same or next line
            '_underline_fills': [],  # No underlines in this template
            # The template has questions as paragraphs - we insert answers after key sections
            'How does the request align with CMS mission and strategic priorities?':
                f"How does the request align with CMS mission and strategic priorities?\n{product} directly supports CMS's digital modernization initiative and reduces manual processing by an estimated 40%.",
            'How do they specifically support our Strategic blueprint OKRs or other CIO/OIT priorities?':
                "How do they specifically support our Strategic blueprint OKRs or other CIO/OIT priorities?\nAligns with OKR 2.3 (Modernize IT Infrastructure) and OKR 4.1 (Improve Operational Efficiency).",
            'What specific problem does the investment solve and the expected impact?':
                f"What specific problem does the investment solve and the expected impact?\nCurrent manual processes result in {random.randint(5,20)}-day turnaround times. {product} reduces this to under 24 hours.",
            'Which stakeholders will it benefit, and how?':
                f"Which stakeholders will it benefit, and how?\nPrimary: {random.choice(['CMCS', 'CCIIO', 'CM', 'OIT'])} staff. Secondary: contractor workforce and external partners.",
            'What is the total cost including development, deployment, and maintenance? How much have we invested to-date?':
                f"What is the total cost including development, deployment, and maintenance? How much have we invested to-date?\nTotal: {self.format_currency(cost)}. Invested to-date: {self.format_currency(cost * 0.3)}.",
        }

    def generate_supplemental_afr_blank_data(self) -> Dict[str, Any]:
        """Generate data for blank Supplemental AFR form (CUI-Financial).

        Fills labeled fields and tables with synthetic budget data.
        """
        PRODUCTS = [
            'FedRAMP Cloud Broker', 'Continuous Diagnostics Platform',
            'Automated Testing Suite', 'Secure Container Runtime',
            'Configuration Management Database', 'IT Asset Discovery Tool',
        ]
        CONTRACTORS = [
            'Accenture Federal', 'Deloitte', 'GDIT', 'Booz Allen Hamilton',
            'CGI Federal', 'Leidos', 'Perspecta', 'SAIC',
        ]
        product = random.choice(PRODUCTS)
        funding = self.generate_currency_amount(200000, 5000000)
        return {
            'Product Description: ': f"Product Description: {product}",
            'Purpose: ': f"Purpose: OIT requires additional funding to procure and deploy {product} to enhance CMS cybersecurity posture and operational efficiency.",
            'Target Users: ': f"Target Users: OIT staff, {random.choice(['CMCS', 'CCIIO', 'CM'])} program teams, and contractor support personnel ({random.randint(50, 500)} users).",
            'Contractor: ': f"Contractor: {random.choice(CONTRACTORS)}",
            'Key Features: ': "Key Features: Automated vulnerability scanning, real-time compliance dashboards, integrated incident response workflows.",
            'Features supported by Funding:': "Features supported by Funding: License renewal, deployment support, and 12-month managed services.",
            'Business Value:': f"Business Value: Reduces mean time to detect threats from {random.randint(24, 72)} hours to under 1 hour. Eliminates {random.randint(100, 500)} hours/year of manual compliance reporting.",
            '_table_data': [
                {
                    'table_index': 0,
                    'start_row': 1,
                    'rows': [[
                        'Operational Efficiency',
                        f"Automates {random.choice(['compliance reporting', 'vulnerability scanning', 'access reviews'])}",
                        f"{random.randint(20, 60)}% reduction in manual effort",
                    ]],
                },
                {
                    'table_index': 1,
                    'start_row': 1,
                    'rows': [[
                        f"{product} License + Support",
                        self.format_currency(funding),
                        self.format_currency(funding * random.uniform(1.5, 3.0)),
                    ]],
                },
            ],
        }

    # --- Legal / FOIA generators ---

    def generate_foia_medicare_auth_data(self) -> Dict[str, Any]:
        """Generate data for FOIA Medicare Records Authorization Form (CUI-Legal).

        Fills 27 PDF form fields: beneficiary info, record details, authorization.
        """
        first = self.fake.first_name()
        last = self.fake.last_name()
        middle = self.fake.first_name()
        dob = self.fake.date_of_birth(minimum_age=30, maximum_age=90)
        # Medicare ID format: 1 letter + 9 digits (MBI format)
        mbi_letter = random.choice('ACDEFGHJKMNPQRTUVWXY')
        mbi = f"{mbi_letter}{random.randint(100000000, 999999999)}"
        start_date = self.fake.date_between(start_date='-5y', end_date='-1y')
        end_date = self.fake.date_between(start_date='-1y', end_date='today')
        states = ['AL','AK','AZ','AR','CA','CO','CT','DE','FL','GA','HI','ID','IL',
                  'IN','IA','KS','KY','LA','ME','MD','MA','MI','MN','MS','MO','MT',
                  'NE','NV','NH','NJ','NM','NY','NC','ND','OH','OK','OR','PA','RI',
                  'SC','SD','TN','TX','UT','VT','VA','WA','WV','WI','WY']

        return {
            'FirstName': first,
            'MiddleName': middle,
            'LastName': last,
            'MedicareID': mbi,
            'Birthdate': dob.strftime('%m/%d/%Y'),
            'StreetAddress': self.fake.street_address(),
            'City': self.fake.city(),
            'State': random.choice(states),
            'Zipcode': self.fake.zipcode(),
            'ReleaseRecords': True,
            'TimeframeStart': start_date.strftime('%m/%d/%Y'),
            'TimeframeEnd': end_date.strftime('%m/%d/%Y'),
            'IndividualRequest': random.choice([True, False]),
            'LitigationRequest': random.choice([True, False]),
            'BeneRecipient': True,
            'Recipient1-Name': self.fake.name(),
            'Recipient1-Email': self.fake.email(),
            'Recipient1-MailingAddress': self.fake.address().replace('\n', ', '),
            'AuthorizationExpiry': True,
            'AuthorizationExpDate': self.fake.date_between(start_date='+30d', end_date='+2y').strftime('%m/%d/%Y'),
            'SignedDate': datetime.now().strftime('%m/%d/%Y'),
            'Signature1': f"{first} {last}",
        }

    def generate_b6_letter_data(self) -> Dict[str, Any]:
        """Generate data for FOIA B6 Letter — partial release response (CUI-Legal).

        Replaces: NAME, COMPANY, ADDRESS, control number, staff names, dates, page counts.
        """
        requester = self.fake.name()
        company = self.fake.company()
        director, liaison = self._foia_staff_names()
        total_pages = random.randint(15, 200)
        released = random.randint(int(total_pages * 0.4), int(total_pages * 0.8))
        withheld = total_pages - released
        request_date = self.fake.date_between(start_date='-1y', end_date='-30d')
        gender = random.choice(['Ms.', 'Mr.'])
        return {
            'NAME': requester,
            'COMPANY': company,
            'ADDRESS': self.fake.street_address(),
            'CITY, STATE ZIP': f"{self.fake.city()}, {self.fake.state_abbr()} {self.fake.zipcode()}",
            '123456789': str(random.randint(100000000, 999999999)),
            'ABCD': ''.join(random.choices('ABCDEFGHIJKLMNOPQRSTUVWXYZ', k=4)),
            'MS/MR.': f'{gender}',
            '6/30/2020': request_date.strftime('%-m/%-d/%Y'),
            'fifty (50)': f"{self._number_word(total_pages)} ({total_pages})",
            'Twenty-five': self._number_word(withheld),
            'Hugh Gilmore': director,
            'Joseph Tripline': liaison,
            'LLC.': f"{company}.",
        }

    def generate_full_release_data(self) -> Dict[str, Any]:
        """Generate data for FOIA Full Release letter (CUI-Legal).

        Replaces: NAME, ADDRESS, EMAIL, requester name, staff names, dates, page counts.
        """
        requester = self.fake.name()
        director, liaison = self._foia_staff_names()
        total_pages = random.randint(5, 100)
        request_date = self.fake.date_between(start_date='-1y', end_date='-30d')
        return {
            'NAME\nADDRESS': f"{requester}\n{self.fake.street_address()}",
            'ADDRESS': f"{self.fake.city()}, {self.fake.state_abbr()} {self.fake.zipcode()}",
            'EMAIL': self.fake.email(),
            'Ms. Halldorsson': f"{'Ms.' if random.random() < 0.5 else 'Mr.'} {requester.split()[-1]}",
            '2/28/2020': request_date.strftime('%-m/%-d/%Y'),
            'twenty-four pages, (NO #S IN PARENS)': f"{self._number_word(total_pages)} pages, ({total_pages})",
            '123456789': str(random.randint(100000000, 999999999)),
            'ABCD': ''.join(random.choices('ABCDEFGHIJKLMNOPQRSTUVWXYZ', k=4)),
            'Hugh Gilmore': director,
            'Joseph Tripline': liaison,
        }

    def generate_form339_letter_data(self) -> Dict[str, Any]:
        """Generate data for FOIA Form 339 Exemption 4 withholding letter (CUI-Legal).

        Replaces: Insert Name, reference number, dates, page counts, staff name.
        """
        requester = self.fake.name()
        director, _ = self._foia_staff_names()
        total_pages = random.randint(10, 150)
        released = random.randint(int(total_pages * 0.3), int(total_pages * 0.7))
        request_date = self.fake.date_between(start_date='-1y', end_date='-30d')
        return {
            'Insert Name': requester,
            'XXXXXXXX': str(random.randint(100000000, 999999999)),
            '(insert date)': request_date.strftime('%B %d, %Y'),
            '(insert recipient)': random.choice([
                'the Centers for Medicare & Medicaid Services',
                'the Department of Health and Human Services',
                'the CMS Freedom of Information Group',
            ]),
            '(insert description of documents, and the CMS 339-Questionnaire)':
                f"{random.randint(2, 8)} CMS-339 Questionnaire forms and supporting financial documentation",
            '(insert number) pages, I have determined to release (insert number)':
                f"{total_pages} pages, I have determined to release {released}",
            'Hugh Gilmore': director,
        }

    def generate_subpoena_response_data(self) -> Dict[str, Any]:
        """Generate data for FOIA Subpoena Duces Tecum response (CUI-Legal).

        Replaces: (INSERT DATE HERE) x2.
        """
        subpoena_date = self.fake.date_between(start_date='-90d', end_date='-7d')
        return {
            '(INSERT DATE HERE)': subpoena_date.strftime('%B %d, %Y'),
        }

    def _number_word(self, n: int) -> str:
        """Convert a number to its word form for page counts."""
        words = {
            1: 'one', 2: 'two', 3: 'three', 4: 'four', 5: 'five',
            6: 'six', 7: 'seven', 8: 'eight', 9: 'nine', 10: 'ten',
            11: 'eleven', 12: 'twelve', 13: 'thirteen', 14: 'fourteen',
            15: 'fifteen', 16: 'sixteen', 17: 'seventeen', 18: 'eighteen',
            19: 'nineteen', 20: 'twenty', 30: 'thirty', 40: 'forty',
            50: 'fifty', 60: 'sixty', 70: 'seventy', 80: 'eighty', 90: 'ninety',
        }
        if n in words:
            return words[n].capitalize()
        if n < 100:
            tens = (n // 10) * 10
            ones = n % 10
            return f"{words.get(tens, str(tens))}-{words.get(ones, str(ones))}".capitalize()
        return str(n)

    def _contract_number(self) -> str:
        return f"75FCMC{random.randint(20, 25)}F{random.randint(1000, 9999)}"

    def _task_order_number(self) -> str:
        return f"75FCMC{random.randint(20, 25)}T{random.randint(10000, 99999)}"

    def generate_currency_amount(self, min_amt: int = 1000, max_amt: int = 10000000) -> float:
        return round(random.uniform(min_amt, max_amt), 2)

    def format_currency(self, amount: float) -> str:
        return f"${amount:,.2f}"

    def generate_rfc_memo_data(self) -> Dict[str, Any]:
        """Generate data for RFC Memo (CUI-Procurement).

        Fills 4 underline blanks: modification checkmarks and signature block.
        """
        return self._rfc_memo_base_data('75FCMCXXXXXXXX / (if applicable) 75FCMCXXXXXXXX')

    def generate_agx_rfc_memo_data(self) -> Dict[str, Any]:
        """Generate data for AGX RFC Memo (CUI-Procurement).

        Same structure as RFC Memo with slightly different contract format.
        """
        return self._rfc_memo_base_data('75FCMC 23F0113 / (if applicable) 75FCMCXXXXXXXX')

    def generate_ja_limited_source_data(self) -> Dict[str, Any]:
        """Generate data for JA Limited Source Justification (CUI-Procurement).

        Fills 9 underline blanks (signature/approval lines).
        """
        names = [self.fake.name() for _ in range(5)]
        dates = [self.fake.date_between(start_date='-90d', end_date='today').strftime('%m/%d/%Y')
                 for _ in range(5)]
        return {
            '_underline_fills': [
                names[0],  # COR signature
                dates[0],
                names[1],  # Program Manager
                dates[1],
                names[2],  # Contracting Officer
                dates[2],
                names[3],  # Competition Advocate
                dates[3],
                f"{names[4]}, Head of Contracting Activity",
            ],
        }

    def generate_jofoc_data(self) -> Dict[str, Any]:
        """Generate data for JOFOC (CUI-Procurement).

        Fills 4 underline blanks and 1 table (option year cost estimates).
        """
        base_cost = self.generate_currency_amount(500000, 5000000)
        return {
            '_underline_fills': [
                random.choice(['Unique Source', 'Unusual and Compelling Urgency',
                               'Statutory Authority', 'National Security']),
                self.fake.name(),  # COR
                self.fake.name(),  # CO
                self.fake.name(),  # Competition Advocate
            ],
            '_table_data': [
                {
                    'table_index': 0,
                    'start_row': 1,
                    'rows': [[
                        self.format_currency(base_cost),
                        self.format_currency(base_cost * 1.03),
                        self.format_currency(base_cost * 1.06),
                        self.format_currency(base_cost * 1.09),
                        self.format_currency(base_cost * 1.12),
                    ]],
                }
            ],
        }

    def generate_oagm_source_selection_data(self) -> Dict[str, Any]:
        """Generate data for OAGM Source Selection Determination (CUI-Procurement).

        Fills header table, offeror evaluation table, and signature line.
        """
        offerors = [self.fake.company() for _ in range(random.randint(3, 5))]
        ratings = ['Outstanding', 'Good', 'Acceptable', 'Marginal']

        eval_rows = []
        for offeror in offerors:
            orig_score = random.randint(70, 98)
            fpr_score = orig_score + random.randint(-3, 5)
            orig_cost = self.generate_currency_amount(1000000, 50000000)
            fpr_cost = orig_cost * random.uniform(0.95, 1.05)
            eval_rows.append([
                offeror,
                f"{orig_score} - {random.choice(ratings)}",
                f"{min(fpr_score, 100)} - {random.choice(ratings[:2])}",
                self.format_currency(orig_cost),
                self.format_currency(fpr_cost),
            ])

        return {
            '_underline_fills': [
                f"{self.fake.name()}    {datetime.now().strftime('%m/%d/%Y')}",
            ],
            '_table_data': [
                {
                    'table_index': 0,
                    'start_row': 0,
                    'rows': [
                        [None, datetime.now().strftime('%B %d, %Y')],
                        [None, self.fake.name()],
                        [None, 'Source Selection Determination'],
                    ],
                },
                {
                    'table_index': 1,
                    'start_row': 1,
                    'rows': eval_rows,
                },
            ],
        }

    def generate_acquisition_plan_data(self) -> Dict[str, Any]:
        """Generate data for CMS Streamlined Acquisition Plan (CUI-Procurement).

        Fills Table 0 (basic info header) and underline blanks.
        """
        project_titles = [
            'Enterprise Cloud Hosting Services',
            'Medicare Beneficiary Data Analytics',
            'Cybersecurity Operations Center Support',
            'Healthcare.gov Platform Maintenance',
            'Quality Payment Program IT Support',
            'Marketplace IT Infrastructure',
        ]
        return {
            '_underline_fills': [
                random.choice(['Mission Critical', 'Business Essential', 'Business Support']),
                self.format_currency(self.generate_currency_amount(1000000, 50000000)),
                f"LCID-{random.randint(2024, 2026)}-{random.randint(100, 999)}",
            ],
            '_table_data': [
                {
                    'table_index': 0,
                    'start_row': 0,
                    'rows': [
                        [None, ''],  # Row 0 header
                        [None, random.choice(self.CMS_OFFICES)],
                        [None, random.choice(project_titles)],
                        [None, self.fake.name()],
                        [None, self.fake.phone_number()],
                        [None, f"{self.fake.first_name().lower()}.{self.fake.last_name().lower()}@cms.hhs.gov"],
                        [None, self.format_currency(self.generate_currency_amount(5000000, 200000000))],
                        [None, self.fake.date_between(start_date='today', end_date='+1y').strftime('%m/%d/%Y')],
                        [None, f"AS-{random.randint(2024, 2026)}-{random.randint(100, 999)}"],
                    ],
                },
            ],
        }

    def generate_market_research_data(self) -> Dict[str, Any]:
        """Generate data for TAB D Market Research Report (CUI-Procurement).

        Fills 4 tables: acquisition team, vendor assessments, business size
        counts, and vendor capability list.
        """
        ROLES = [
            'Requirements development and technical evaluation',
            'Acquisition planning and contract execution',
            'Contract administration and oversight',
            'Technical oversight and acceptance',
            'Small business coordination and outreach',
            'Subject matter expertise',
        ]
        CAPABILITIES = [
            'Demonstrated capability in cloud migration and hosting for federal agencies',
            'Strong past performance on similar HHS/CMS contracts',
            'Certified FedRAMP High cloud service provider',
            'Experienced in Agile development with DevSecOps practices',
            'Holds relevant CMMI Level 3 or higher certification',
            'Proven 508 compliance and accessibility testing capabilities',
            'Extensive experience with Medicare/Medicaid IT systems',
            'ISO 27001 certified information security management',
        ]
        BIZ_SIZES = ['Small Business', 'Small Business', 'Large Business',
                     'Small Business', 'Large Business']

        # Table 0: Acquisition team (rows 1-6, col 0=name, 2=office, 3=phone, 4=email)
        team_rows = []
        for ri in range(6):
            name = self.fake.name()
            office = random.choice(self.CMS_OFFICES)
            phone = self.fake.phone_number()
            email = f"{name.split()[0].lower()}.{name.split()[-1].lower()}@cms.hhs.gov"
            role = ROLES[ri] if ri < len(ROLES) else random.choice(ROLES)
            # Only fill cols 0, 2, 3, 4, 5 — col 1 (title) is pre-filled
            team_rows.append([name, None, office, phone, email, role])

        # Table 1: Vendor assessments (3 rows)
        vendor_rows = []
        for _ in range(3):
            vendor_rows.append([
                self.fake.company(),
                f"{self.fake.city()}, {self.fake.state_abbr()}",
                self.fake.name(),
                random.choice(CAPABILITIES),
            ])

        # Table 2: Business size counts (rows 1-10, col 1 = count)
        size_rows = []
        for _ in range(10):
            count = random.choice([0, 0, 1, 2, 3, 5, 8])
            size_rows.append([None, str(count)])

        # Table 3: Vendor capability list (5 rows)
        cap_rows = []
        for _ in range(5):
            cap_rows.append([
                self.fake.company(),
                random.choice(BIZ_SIZES),
                random.choice(CAPABILITIES),
            ])

        return {
            '_table_data': [
                {'table_index': 0, 'start_row': 1, 'rows': team_rows},
                {'table_index': 1, 'start_row': 1, 'rows': vendor_rows},
                {'table_index': 2, 'start_row': 1, 'rows': size_rows},
                {'table_index': 3, 'start_row': 1, 'rows': cap_rows},
            ]
        }

    def generate_clin_template_data(self) -> Dict[str, Any]:
        """Generate data for CLIN Templates (CUI-Procurement).

        Fills all 3 contract line item tables (Fixed Price, Cost Reimbursement, T&M)
        with 2-4 synthetic CLINs each.
        """
        PSC_CODES = [
            'D302', 'D306', 'D307', 'D308', 'D310', 'D311', 'D314', 'D316', 'D317', 'D399',
            'R408', 'R413', 'R425', 'R497', 'R499', 'R707', 'R799',
        ]
        CLIN_DESCRIPTIONS = [
            'Cloud Infrastructure Hosting Services',
            'Cybersecurity Operations and Monitoring',
            'Application Development and Maintenance',
            'Enterprise Data Analytics Platform',
            'Help Desk and End User Support',
            'IT Program Management Support',
            'Network Operations Center Services',
            'Identity and Access Management Services',
            'System Integration and Testing',
            'Cloud Migration and Modernization',
            'Database Administration Services',
            'Quality Assurance and IV&V',
            'Disaster Recovery and COOP Services',
            'Section 508 Compliance Testing',
            'DevSecOps Pipeline Support',
        ]
        UNITS = ['Month', 'Each', 'Hour', 'Lot', 'Year']

        def _acct_class():
            return f"75-{random.randint(100,999)}0-0-1-{random.randint(300,399)}"

        def _make_ffp_rows(n):
            rows = []
            for i in range(n):
                clin = f"000{i+1}"
                desc = random.choice(CLIN_DESCRIPTIONS)
                unit = random.choice(UNITS)
                qty = random.randint(1, 24)
                price = round(random.uniform(5000, 150000), 2)
                total = round(price * qty, 2)
                rows.append([clin, desc, random.choice(PSC_CODES), _acct_class(),
                             unit, str(qty), f"${price:,.2f}", f"${total:,.2f}"])
            return rows

        def _make_cr_rows(n):
            rows = []
            for i in range(n):
                clin = f"000{i+1}"
                desc = random.choice(CLIN_DESCRIPTIONS)
                unit = random.choice(UNITS)
                qty = random.randint(1, 12)
                cost = round(random.uniform(50000, 500000), 2)
                fee = round(cost * random.uniform(0.05, 0.10), 2)
                total = round(cost + fee, 2)
                pop_start = self.fake.date_between(start_date='-1y', end_date='today')
                pop_end = self.fake.date_between(start_date='today', end_date='+2y')
                pop = f"{pop_start.strftime('%m/%d/%Y')} - {pop_end.strftime('%m/%d/%Y')}"
                rows.append([clin, desc, random.choice(PSC_CODES), _acct_class(),
                             unit, str(qty), f"${cost:,.2f}", f"${fee:,.2f}",
                             f"${total:,.2f}", pop])
            return rows

        def _make_tm_rows(n):
            rows = []
            for i in range(n):
                clin = f"000{i+1}"
                desc = random.choice(CLIN_DESCRIPTIONS)
                unit = random.choice(['Hour', 'Month'])
                qty = random.randint(100, 5000) if unit == 'Hour' else random.randint(6, 24)
                cost = round(random.uniform(75000, 750000), 2)
                pop_start = self.fake.date_between(start_date='-1y', end_date='today')
                pop_end = self.fake.date_between(start_date='today', end_date='+2y')
                pop = f"{pop_start.strftime('%m/%d/%Y')} - {pop_end.strftime('%m/%d/%Y')}"
                rows.append([clin, desc, random.choice(PSC_CODES), _acct_class(),
                             unit, str(qty), f"${cost:,.2f}", pop])
            return rows

        num_clins = random.randint(2, 4)
        return {
            '_table_data': [
                {'table_index': 0, 'start_row': 3, 'rows': _make_ffp_rows(num_clins)},
                {'table_index': 1, 'start_row': 3, 'rows': _make_cr_rows(num_clins)},
                {'table_index': 2, 'start_row': 3, 'rows': _make_tm_rows(num_clins)},
            ]
        }

    def generate_reasonable_accommodation_data(self) -> Dict[str, Any]:
        """Generate data for Reasonable Accommodation Request (CUI)."""

        employee_name = self.fake.name()

        accommodations = [
            "Modified work schedule to accommodate medical appointments",
            "Ergonomic keyboard and mouse for repetitive strain injury",
            "Screen reader software for visual impairment",
            "Reserved parking space near building entrance",
            "Standing desk for back condition",
            "Noise-canceling headphones for concentration",
            "Remote work option for chronic condition management",
        ]

        form_data = {
            'Name': employee_name,  # Actual field name in PDF
            'Date of Birth': self.fake.date_of_birth(minimum_age=25, maximum_age=65).strftime('%m/%d/%Y'),
            'Grade': random.choice(['GS-9', 'GS-11', 'GS-12', 'GS-13', 'GS-14', 'GS-15']),
            'Component': random.choice(['CMS', 'OIG', 'ACF', 'ASPE', 'OCR']),
            'Location': self.fake.city() + ', ' + self.fake.state_abbr(),
            'Telephone number': self.fake.phone_number(),
            'Manager': self.fake.name(),
            'Discription': random.choice(accommodations),  # Note: typo in actual PDF field name
        }

        return form_data


class CustomerTemplateManager:
    """Manages customer-provided template forms."""

    def __init__(self, template_dir: str = 'temp', output_dir: str = 'output'):
        """
        Initialize template manager.

        Args:
            template_dir: Directory containing customer template files
            output_dir: Base output directory
        """
        self.template_dir = template_dir
        self.output_dir = output_dir
        self.populator = PDFFormPopulator()

        # Map templates to data generators
        self.template_mappings = {
            'Medical Inquiry  Form': {
                'template': 'saved_templates/Medical Inquiry  Form_508-blank-PHI-negative.pdf',
                'generator': self.populator.generate_medical_inquiry_data,
                'category': 'PHI',
                'clean_name': 'MedicalInquiryForm',
                'acroform': True,
            },
            'EFT Authorization Form': {
                'template': 'saved_templates/EFT Authorization Form-blank-CUI-Finance-negative.pdf',
                'generator': self.populator.generate_eft_authorization_data,
                'category': 'CUI-Finance',
                'clean_name': 'EFTAuthorizationForm',
                'acroform': True,
            },
            'ReasonableAccommodationRequest': {
                'template': 'saved_templates/ReasonableAccommodationRequest-blank-CUI-negative.pdf',
                'generator': self.populator.generate_reasonable_accommodation_data,
                'category': 'CUI',
                'clean_name': 'ReasonableAccommodationRequest',
                'acroform': True,
            },
            # Procurement: IGCE XLSX (positive-only, copy mode — openpyxl can't parse drawings)
            'IGCE': {
                'template_positive': 'IaaS Mainframe MFA IGCE OY1-CUI-Procurement and Acquisition-positive.xlsx',
                'category': 'CUI-Procurement',
                'clean_name': 'IGCE',
                'positive_only': True,
            },
            'CLINTemplates': {
                'template': 'CLIN Templates-CUI-Procurement-positive.docx',
                'generator': self.populator.generate_clin_template_data,
                'category': 'CUI-Procurement',
                'clean_name': 'CLINTemplates',
                'positive_only': True,
            },
            'MarketResearch': {
                'template': 'TAB D Market Research-CUI-Procurement-positive.docx',
                'generator': self.populator.generate_market_research_data,
                'category': 'CUI-Procurement',
                'clean_name': 'MarketResearch',
                'positive_only': True,
            },
            'RFCMemo': {
                'template': 'RFC Memo-CUI-Procurement-positive.docx',
                'generator': self.populator.generate_rfc_memo_data,
                'category': 'CUI-Procurement',
                'clean_name': 'RFCMemo',
                'positive_only': True,
            },
            'AGXRFCMemo': {
                'template': 'AGX RFC Memo-CUI-Procurement-positive.docx',
                'generator': self.populator.generate_agx_rfc_memo_data,
                'category': 'CUI-Procurement',
                'clean_name': 'AGXRFCMemo',
                'positive_only': True,
            },
            'JALimitedSource': {
                'template': 'JA LimitedSource-CUI-Procurement-positive.docx',
                'generator': self.populator.generate_ja_limited_source_data,
                'category': 'CUI-Procurement',
                'clean_name': 'JALimitedSource',
                'positive_only': True,
            },
            'JOFOC': {
                'template': 'JOFOC-CUI-Procurement-positive.docx',
                'generator': self.populator.generate_jofoc_data,
                'category': 'CUI-Procurement',
                'clean_name': 'JOFOC',
                'positive_only': True,
            },
            'OAGMSourceSelection': {
                'template': 'OAGM SourceSelection-CUI-Procurement-positive.docx',
                'generator': self.populator.generate_oagm_source_selection_data,
                'category': 'CUI-Procurement',
                'clean_name': 'OAGMSourceSelection',
                'positive_only': True,
            },
            'AcquisitionPlan': {
                'template': 'CMS AcquisitionPlan-CUI-Procurement-positive.docx',
                'generator': self.populator.generate_acquisition_plan_data,
                'category': 'CUI-Procurement',
                'clean_name': 'AcquisitionPlan',
                'positive_only': True,
            },
            # Critical Infrastructure: fillable DOCX pairs
            'KMP': {
                'template': 'KMP-MockSystem-CUI-Critical Infrastructure-Positive.docx',
                'template_negative': 'saved_templates/KMPTemplate-CUI-Critical Infrastructure-negative.docx',
                'generator': self.populator.generate_kmp_data,
                'category': 'CUI-CritInfra',
                'clean_name': 'KMP',
            },
            'RulesOfBehavior': {
                'template': '2025 PQCRA Rules of Behavior - MAC NAME-CUI-Critical Infrastructure-positive.docx',
                'template_negative': 'saved_templates/2025 PQCRA Rules of Behavior - MAC NAME-CUI-Critical Infrastructure-negative.docx',
                'generator': self.populator.generate_rules_of_behavior_data,
                'category': 'CUI-CritInfra',
                'clean_name': 'RulesOfBehavior',
            },
            'IncidentResponse': {
                'template': 'saved_templates/rmh-chapter-08-incident-response-incident-report-template-CUI-Critical Infrastructure-positive.docx',
                'template_negative': 'saved_templates/rmh-chapter-08-incident-response-incident-report-template-CUI-Critical Infrastructure-negative.docx',
                'generator': self.populator.generate_incident_response_data,
                'category': 'CUI-CritInfra',
                'clean_name': 'IncidentResponse',
            },
            # Financial: fillable DOCX pairs
            'AFRAdditionalInfo': {
                'template': 'Additional Information OIT FO-Mock AFR-CUI-Budget-positive.docx',
                'template_negative': 'saved_templates/Additional Information OIT FO form-CUI-Budget-negative.docx',
                'generator': self.populator.generate_afr_additional_info_data,
                'category': 'CUI-Financial',
                'clean_name': 'AFRAdditionalInfo',
            },
            'DIBOAFR': {
                'template': 'DIBO AFR -AMI-CUI-Budget-Positive.docx',
                'template_negative': 'saved_templates/DIBO AFR Guidance-template-CUI-Budget-Negative.docx',
                'generator': self.populator.generate_dibo_afr_data,
                'category': 'CUI-Financial',
                'clean_name': 'DIBOAFR',
            },
            'SupplementalAFR': {
                'template': 'Supplemental AFR Information-MockProject-CUI-Budget-Positive.docx',
                'template_negative': 'saved_templates/Supplemental AFR Information-template-CUI-Budget-negative.docx',
                'generator': self.populator.generate_supplemental_afr_data,
                'category': 'CUI-Financial',
                'clean_name': 'SupplementalAFR',
            },
            # Financial: additional fillable templates
            'OITFOAdditionalInfo': {
                'template': 'saved_templates/OIT FO - AIF-CUI-Budget-negative.docx',
                'generator': self.populator.generate_oit_fo_aif_data,
                'category': 'CUI-Financial',
                'clean_name': 'OITFOAdditionalInfo',
                'positive_only': True,
            },
            'SupplementalAFRBlank': {
                'template': 'saved_templates/Supplemental AFR-CUI-Budget-negative.docx',
                'generator': self.populator.generate_supplemental_afr_blank_data,
                'category': 'CUI-Financial',
                'clean_name': 'SupplementalAFRBlank',
                'positive_only': True,
            },
            # Critical Infrastructure: HHS RBD fillable PDF (pos+neg pair)
            'HHSRBD': {
                'template': 'hhs-RBD-CUI-Critical Infrastructure-positive.pdf',
                'template_negative': 'saved_templates/hhs-RBD-CUI-Critical Infrastructure-negative.pdf',
                'generator': self.populator.generate_hhs_rbd_data,
                'category': 'CUI-CritInfra',
                'clean_name': 'HHSRBD',
                'acroform': True,
            },
            # Legal / FOIA: fillable positives with negative pairs
            'B6Letter': {
                'template': 'B6 Letter-CUI-Legal-positive.docx',
                'template_negative': 'saved_templates/FOIA RequestLetter-CUI-Legal-negative.pdf',
                'generator': self.populator.generate_b6_letter_data,
                'category': 'CUI-Legal',
                'clean_name': 'B6Letter',
            },
            'FullRelease': {
                'template': 'Full Release-CUI-Legal-positive.docx',
                'template_negative': 'saved_templates/FOIA RequestLetterLivingBene-CUI-Legal-negative.pdf',
                'generator': self.populator.generate_full_release_data,
                'category': 'CUI-Legal',
                'clean_name': 'FullRelease',
            },
            'Form339Letter': {
                'template': 'Form339 Letter-CUI-Legal-positive.docx',
                'template_negative': 'saved_templates/FOIA RequestLetterOwnRecords-CUI-Legal-negative.pdf',
                'generator': self.populator.generate_form339_letter_data,
                'category': 'CUI-Legal',
                'clean_name': 'Form339Letter',
            },
            'SubpoenaResponse': {
                'template': 'Subpoena Response-CUI-Legal-positive.docx',
                'template_negative': 'saved_templates/FOIA Appeal-CUI-Legal-negative.pdf',
                'generator': self.populator.generate_subpoena_response_data,
                'category': 'CUI-Legal',
                'clean_name': 'SubpoenaResponse',
            },
            # Legal / FOIA: fillable positive only (Medicare Auth — contains PHI-level data)
            'FOIAMedicareAuth': {
                'template': 'saved_templates/FOIA MedicareAuth-CUI-Legal-positive.pdf',
                'generator': self.populator.generate_foia_medicare_auth_data,
                'category': 'CUI-Legal',
                'clean_name': 'FOIAMedicareAuth',
                'positive_only': True,
                'acroform': True,
            },
            # Legal / FOIA: copy-only negatives
            'FOIAGuidance': {
                'template_positive': 'saved_templates/FOIA Guidance-CUI-Legal-negative.docx',
                'category': 'CUI-Legal',
                'clean_name': 'FOIAGuidance',
                'negative_only': True,
            },
            'FOIARequestDeceasedBene': {
                'template_positive': 'saved_templates/FOIA RequestLetterDeceasedBene-CUI-Legal-negative.pdf',
                'category': 'CUI-Legal',
                'clean_name': 'FOIARequestDeceasedBene',
                'negative_only': True,
            },
        }

    def generate_from_template(self, template_key: str, output_subdir: str,
                               index: int, populate: bool = True,
                               extra_data: dict = None,
                               field_data: dict = None) -> str:
        """
        Generate a document from customer template.

        Args:
            template_key: Key from template_mappings
            output_subdir: Full path to output directory (not relative)
            index: Document index for filename
            populate: If True, populate with data. If False, use blank template.
            extra_data: Optional dict of additional fields to merge (e.g., LLM narratives)

        Returns:
            Path to generated file
        """
        template_info = self.template_mappings[template_key]

        # Skip positive-only templates when generating negatives
        if not populate and template_info.get('positive_only'):
            return None
        # Skip negative-only templates when generating positives
        if populate and template_info.get('negative_only'):
            return None

        # Choose template based on positive/negative and whether we have separate templates
        if 'template_positive' in template_info:
            # Has separate positive/negative templates (e.g., EFT)
            if populate:
                template_file = template_info['template_positive']
            elif 'template_negative' in template_info:
                template_file = template_info['template_negative']
            elif template_info.get('negative_only'):
                # negative_only: the template_positive file IS the negative content
                template_file = template_info['template_positive']
            else:
                return None  # No negative template available
            # Just copy the appropriate template (positive already has data)
            template_path = os.path.join(self.template_dir, template_file)
            clean_name = template_info['clean_name']
            ext = os.path.splitext(template_file)[1]  # Detect extension from template
            filename = f"{clean_name}_{index:04d}{ext}"
            output_path = os.path.join(output_subdir, filename)

            import shutil
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            shutil.copy(template_path, output_path)
            return output_path
        else:
            # Single template - need to populate
            template_path = os.path.join(self.template_dir, template_info['template'])
            clean_name = template_info['clean_name']
            ext = os.path.splitext(template_info['template'])[1]  # Detect extension from template
            filename = f"{clean_name}_{index:04d}{ext}"
            output_path = os.path.join(output_subdir, filename)

            if populate:
                # Use pre-generated data if provided, otherwise generate now
                if field_data is None:
                    field_data = template_info['generator']()
                # Merge LLM-enriched narratives if provided
                if extra_data:
                    field_data.update(extra_data)
                if ext.lower() == '.pdf':
                    if template_info.get('acroform'):
                        return self.populator.populate_acroform(
                            template_path, output_path, field_data)
                    return self.populator.populate_form(
                        template_path, output_path, field_data,
                        field_positions=template_info.get('field_positions'))
                elif ext.lower() == '.docx':
                    return self.populate_docx_template(template_path, output_path, field_data)
                else:
                    # For other formats, copy and let caller handle
                    import shutil
                    os.makedirs(os.path.dirname(output_path), exist_ok=True)
                    shutil.copy(template_path, output_path)
                    return output_path
            else:
                # Copy blank/negative template
                import shutil
                # Use separate negative template if available
                if 'template_negative' in template_info:
                    neg_path = os.path.join(self.template_dir, template_info['template_negative'])
                    ext = os.path.splitext(template_info['template_negative'])[1]
                    filename = f"{clean_name}_{index:04d}{ext}"
                    output_path = os.path.join(output_subdir, filename)
                    template_path = neg_path
                os.makedirs(os.path.dirname(output_path), exist_ok=True)
                shutil.copy(template_path, output_path)
                return output_path

    def populate_docx_template(self, template_path: str, output_path: str,
                                replacements: Dict[str, str]) -> str:
        """
        Populate a DOCX template with synthetic values.

        Supports three fill mechanisms (all optional, can be combined):
        1. Text placeholder substitution: {'PlaceholderText': 'replacement'}
        2. Underline blank fills: {'_underline_fills': ['val1', 'val2', ...]}
           Replaces ___ blanks in document order
        3. Table data fills: {'_table_data': [{'table_index': 0, 'start_row': 3,
           'rows': [['a','b'], ['c','d']]}]}
           Writes values into table cells by position

        Args:
            template_path: Path to DOCX template
            output_path: Path to save populated DOCX
            replacements: Dict with text replacements and optional _table_data/_underline_fills

        Returns:
            Path to created file
        """
        from docx import Document
        import re

        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc = Document(template_path)

        # Extract special keys before text substitution
        table_data = replacements.pop('_table_data', None)
        underline_fills = replacements.pop('_underline_fills', None)
        append_paragraphs = replacements.pop('_append_paragraphs', None)
        # NOTE: _underline_fills values are consumed in iteration order:
        # body paragraphs → table cells → headers/footers.
        # Generators must provide values matching this order, not visual page order.
        # If a template mixes underlines across these regions, use _table_data instead.
        underline_iter = iter(underline_fills) if underline_fills else None

        def replace_in_runs(paragraph):
            """Replace placeholder text in paragraph runs, preserving formatting."""
            full_text = paragraph.text
            changed = False

            # Standard placeholder substitution
            for placeholder, value in replacements.items():
                if placeholder in full_text:
                    full_text = full_text.replace(placeholder, str(value))
                    changed = True

            # Underline blank substitution: replace each ___ with next value
            if underline_iter and re.search(r'_{3,}', full_text):
                def replace_blank(match):
                    try:
                        return next(underline_iter)
                    except StopIteration:
                        return match.group(0)
                full_text = re.sub(r'_{3,}', replace_blank, full_text)
                changed = True

            if changed and paragraph.runs:
                paragraph.runs[0].text = full_text
                for run in paragraph.runs[1:]:
                    run.text = ""

        for paragraph in doc.paragraphs:
            replace_in_runs(paragraph)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_runs(paragraph)

        # Also check headers and footers
        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header and header.is_linked_to_previous is False:
                    for paragraph in header.paragraphs:
                        replace_in_runs(paragraph)
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer and footer.is_linked_to_previous is False:
                    for paragraph in footer.paragraphs:
                        replace_in_runs(paragraph)

        # Table data fill: write values into empty cells by position
        if table_data:
            self._populate_docx_tables(doc, table_data)

        # Append LLM narrative sections at end of document
        if append_paragraphs:
            from docx.shared import Pt
            doc.add_paragraph()  # Spacer
            for section in append_paragraphs:
                heading = section.get('heading', '')
                text = section.get('text', '')
                if heading:
                    h = doc.add_paragraph()
                    run = h.add_run(heading)
                    run.bold = True
                    run.font.size = Pt(11)
                if text:
                    for para in text.split('\n\n'):
                        if para.strip():
                            doc.add_paragraph(para.strip())

        doc.save(output_path)
        return output_path

    def _populate_docx_tables(self, doc, table_data: list):
        """
        Fill table rows with synthetic data.

        Args:
            doc: python-docx Document (mutated in place)
            table_data: list of dicts:
                {
                    'table_index': int,   # which table in the doc
                    'start_row': int,     # first data row to fill
                    'rows': [             # data rows
                        ['val1', 'val2', ...],
                    ]
                }
        """
        for entry in table_data:
            table_idx = entry['table_index']
            start_row = entry['start_row']
            rows = entry['rows']

            if table_idx >= len(doc.tables):
                continue

            table = doc.tables[table_idx]
            for row_offset, row_values in enumerate(rows):
                row_idx = start_row + row_offset
                if row_idx >= len(table.rows):
                    break
                row = table.rows[row_idx]
                for col_idx, value in enumerate(row_values):
                    if col_idx >= len(row.cells) or value is None:
                        continue  # None = skip, preserve existing cell content
                    cell = row.cells[col_idx]
                    if cell.paragraphs:
                        cell.paragraphs[0].text = str(value)
                    else:
                        cell.text = str(value)

    def list_available_templates(self):
        """List all available customer templates."""
        print("\nAvailable Customer Templates:")
        print("="*70)
        for key, info in self.template_mappings.items():
            print(f"\n{key}")
            print(f"  Template: {info['template']}")
            print(f"  Category: {info['category']}")
            print(f"  Output Name: {info['clean_name']}")
        print("\n" + "="*70)


if __name__ == "__main__":
    # Test the populator
    manager = CustomerTemplateManager()
    manager.list_available_templates()

    # Test generating a populated Medical Inquiry Form
    print("\nTesting Medical Inquiry Form population...")
    output_file = manager.generate_from_template(
        'Medical Inquiry  Form',
        'temp',
        1,
        populate=True
    )
    print(f"Created: {output_file}")

    # Test blank version
    print("\nTesting blank form...")
    blank_file = manager.generate_from_template(
        'Medical Inquiry  Form',
        'temp',
        2,
        populate=False
    )
    print(f"Created: {blank_file}")
