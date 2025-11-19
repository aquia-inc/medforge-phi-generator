"""
Enhanced DOCX formatter with LLM-generated clinical narratives
Uses Claude 4.5 Sonnet for 20% of documents
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os
import random
from generators.llm_generator import ClaudeGenerator, is_llm_available


class EnhancedPHIDocxFormatter:
    """Creates DOCX documents with optional LLM-enhanced narratives"""

    def __init__(self, output_dir='output', llm_percentage=0.2, use_llm=True):
        """
        Initialize formatter

        Args:
            output_dir: Output directory for documents
            llm_percentage: Percentage of documents to use LLM (0.0-1.0)
            use_llm: Whether to attempt LLM usage (falls back to templates if unavailable)
        """
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        self.llm_percentage = llm_percentage
        self.use_llm = use_llm and is_llm_available()

        if self.use_llm:
            try:
                self.claude_gen = ClaudeGenerator()
                print(f"✓ LLM enabled: {int(llm_percentage * 100)}% of documents will use Claude-generated narratives")
            except Exception as e:
                print(f"⚠ LLM initialization failed: {e}. Using template-based generation.")
                self.use_llm = False
                self.claude_gen = None
        else:
            self.claude_gen = None
            print("ℹ LLM disabled: Using template-based generation only")

    def _should_use_llm(self) -> bool:
        """Determine if this document should use LLM based on percentage"""
        return self.use_llm and random.random() < self.llm_percentage

    def create_progress_note_enhanced(self, patient, provider, facility, filename):
        """
        Generate progress note with optional LLM-enhanced narrative

        Args:
            patient: Patient data dict
            provider: Provider data dict
            facility: Facility data dict
            filename: Output filename

        Returns:
            Tuple of (filepath, used_llm: bool)
        """
        doc = Document()
        used_llm = False

        # Header
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header.add_run(facility['name'].upper())
        run.bold = True
        run.font.size = Pt(14)

        doc.add_paragraph()

        # Document title
        title = doc.add_paragraph()
        title_run = title.add_run('PROGRESS NOTE / SOAP NOTE')
        title_run.bold = True
        title_run.font.size = Pt(12)

        doc.add_paragraph()

        # Patient header
        patient_header = doc.add_paragraph()
        patient_header.add_run(
            f"Patient: {patient['last_name']}, {patient['first_name']}  |  "
            f"DOB: {patient['dob'].strftime('%m/%d/%Y')}  |  "
            f"MRN: {patient['mrn']}\n"
            f"Date of Visit: {datetime.now().strftime('%m/%d/%Y')}\n"
            f"Provider: {provider['first_name']} {provider['last_name']}, {provider['title']} - {provider['specialty']}"
        ).font.size = Pt(10)

        doc.add_paragraph()

        # Vital Signs
        doc.add_paragraph().add_run('VITAL SIGNS').bold = True
        vitals = patient['vital_signs']
        vitals_text = (
            f"BP: {vitals['blood_pressure']} mmHg  |  "
            f"HR: {vitals['heart_rate']} bpm  |  "
            f"Temp: {vitals['temperature']}°F  |  "
            f"RR: {vitals['respiratory_rate']}  |  "
            f"O2 Sat: {vitals['oxygen_saturation']}%\n"
            f"Weight: {vitals['weight']} lbs  |  "
            f"Height: {vitals['height']} inches"
        )
        doc.add_paragraph(vitals_text)

        doc.add_paragraph()

        # SOAP Format - Try LLM first
        if self._should_use_llm() and self.claude_gen:
            try:
                narrative = self.claude_gen.generate_clinical_narrative(
                    patient=patient,
                    diagnoses=patient['diagnoses'],
                    medications=patient['medications'],
                    vitals=patient['vital_signs']
                )

                # Subjective
                doc.add_paragraph().add_run('SUBJECTIVE:').bold = True
                doc.add_paragraph(narrative.subjective)

                # Objective
                doc.add_paragraph().add_run('OBJECTIVE:').bold = True
                doc.add_paragraph(narrative.objective)

                # Assessment
                doc.add_paragraph().add_run('ASSESSMENT:').bold = True
                doc.add_paragraph(narrative.assessment)

                # Plan
                doc.add_paragraph().add_run('PLAN:').bold = True
                doc.add_paragraph(narrative.plan)

                used_llm = True

            except Exception as e:
                print(f"⚠ LLM generation failed, using template: {e}")
                self._add_template_soap(doc, patient)
        else:
            # Template-based SOAP
            self._add_template_soap(doc, patient)

        doc.add_paragraph()

        # Signature
        sig = doc.add_paragraph()
        sig.add_run(
            f"\n---\n"
            f"Electronically signed by: {provider['first_name']} {provider['last_name']}, {provider['title']}\n"
            f"Date: {datetime.now().strftime('%m/%d/%Y %H:%M')}\n"
            f"NPI: {provider['npi']}"
        ).font.size = Pt(9)

        # Save
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        return filepath, used_llm

    def _add_template_soap(self, doc, patient):
        """Add template-based SOAP note sections"""
        # Subjective
        doc.add_paragraph().add_run('SUBJECTIVE:').bold = True
        doc.add_paragraph(
            f"Patient presents for follow-up visit. Reports feeling generally well. "
            f"Current medications include: {', '.join(patient['medications'][:3])}. "
            f"Known allergies: {', '.join(patient['allergies'])}."
        )

        # Objective
        doc.add_paragraph().add_run('OBJECTIVE:').bold = True
        doc.add_paragraph(
            "Physical exam reveals patient in no acute distress. "
            "Vital signs as noted above. "
            "Review of systems within normal limits for age."
        )

        # Assessment
        doc.add_paragraph().add_run('ASSESSMENT:').bold = True
        for diagnosis in patient['diagnoses']:
            doc.add_paragraph(
                f"• {diagnosis['name']} (ICD-10: {diagnosis['icd10']})",
                style='List Bullet'
            )

        # Plan
        doc.add_paragraph().add_run('PLAN:').bold = True
        doc.add_paragraph(
            "1. Continue current medications as prescribed\n"
            "2. Follow-up lab work in 3 months\n"
            "3. Return to clinic in 6 months or sooner if symptoms worsen\n"
            "4. Patient education provided regarding disease management"
        )

    def create_lab_result(self, patient, provider, facility, lab_data, filename):
        """
        Generate lab result document (template-based, no LLM needed for structured data)

        Returns:
            filepath
        """
        doc = Document()

        # Add facility header
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header.add_run(facility['name'].upper())
        run.bold = True
        run.font.size = Pt(16)

        # Facility address
        addr_para = doc.add_paragraph()
        addr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        addr_para.add_run(
            f"{facility['address']['street']}\n"
            f"{facility['address']['city']}, {facility['address']['state']} {facility['address']['zip']}\n"
            f"Phone: {facility['phone']} | Fax: {facility['fax']}"
        ).font.size = Pt(10)

        doc.add_paragraph()

        # Document title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run('LABORATORY RESULTS')
        title_run.bold = True
        title_run.font.size = Pt(14)

        doc.add_paragraph()

        # Patient Information Section
        doc.add_paragraph().add_run('PATIENT INFORMATION').bold = True
        patient_info = doc.add_table(rows=6, cols=2)
        patient_info.style = 'Light Grid Accent 1'

        cells = [
            ('Patient Name:', f"{patient['last_name']}, {patient['first_name']}"),
            ('Date of Birth:', patient['dob'].strftime('%m/%d/%Y')),
            ('Age:', str(patient['age'])),
            ('MRN:', patient['mrn']),
            ('Address:', f"{patient['address']['street']}, {patient['address']['city']}, {patient['address']['state']} {patient['address']['zip']}"),
            ('Phone:', patient['phone'])
        ]

        for i, (label, value) in enumerate(cells):
            patient_info.rows[i].cells[0].text = label
            patient_info.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            patient_info.rows[i].cells[1].text = value

        doc.add_paragraph()

        # Test Information
        doc.add_paragraph().add_run('TEST INFORMATION').bold = True
        test_info = doc.add_table(rows=3, cols=2)
        test_info.style = 'Light Grid Accent 1'

        test_cells = [
            ('Collection Date:', lab_data['test_date'].strftime('%m/%d/%Y')),
            ('Report Date:', datetime.now().strftime('%m/%d/%Y')),
            ('Ordering Provider:', f"{provider['first_name']} {provider['last_name']}, {provider['title']}")
        ]

        for i, (label, value) in enumerate(test_cells):
            test_info.rows[i].cells[0].text = label
            test_info.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            test_info.rows[i].cells[1].text = value

        doc.add_paragraph()

        # Lab Results Table
        doc.add_paragraph().add_run('LABORATORY RESULTS').bold = True
        results_table = doc.add_table(rows=len(lab_data['results']) + 1, cols=5)
        results_table.style = 'Light Grid Accent 1'

        # Headers
        headers = ['Test Name', 'Result', 'Unit', 'Reference Range', 'Flag']
        for i, header in enumerate(headers):
            cell = results_table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True

        # Fill results
        for i, result in enumerate(lab_data['results'], 1):
            results_table.rows[i].cells[0].text = result['test']
            results_table.rows[i].cells[1].text = str(result['value'])
            results_table.rows[i].cells[2].text = result['unit']
            results_table.rows[i].cells[3].text = result['reference_range']
            flag_cell = results_table.rows[i].cells[4]
            flag_cell.text = result.get('flag', '')
            if result.get('flag'):
                flag_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
                flag_cell.paragraphs[0].runs[0].bold = True

        doc.add_paragraph()

        # Footer
        footer = doc.add_paragraph()
        footer.add_run(
            '---\n'
            'This report contains confidential patient health information. '
            'Distribution or copying is prohibited without authorization.\n'
            f'Medical Director: {provider["first_name"]} {provider["last_name"]}, {provider["title"]}\n'
            f'NPI: {provider["npi"]}'
        ).font.size = Pt(8)

        # Save
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        return filepath
