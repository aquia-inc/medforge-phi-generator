"""
CUI nested email formatter.

Creates CUI emails with document attachments (PDF, DOCX, or ZIP),
mirroring the PHI NestedEmailFormatter capabilities.
"""
import io
import random
import zipfile
from typing import Any, Dict

from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from docx import Document

from formatters.base_email_formatter import BaseEmailFormatter


class CUINestedEmailFormatter(BaseEmailFormatter):
    """Creates CUI emails with document attachments."""

    def create_cui_email_with_attachment(self, doc_data: Dict[str, Any], filename: str,
                                        is_positive: bool = True) -> str:
        """
        Create CUI email with an embedded attachment (PDF, DOCX, or ZIP).

        80% single PDF/DOCX, 20% ZIP with multiple files.

        Args:
            doc_data: CUI document data dict
            filename: Output filename
            is_positive: True for CUI-positive, False for CUI-negative

        Returns:
            Path to created EML file
        """
        classification = doc_data.get('classification', '')
        has_cui = doc_data.get('has_cui', False)
        title = doc_data.get('title', 'Document')
        agency = doc_data.get('agency', 'Department of Health and Human Services')
        agency_domain = agency.lower().replace(' ', '').replace('of', '')[:10] + '.gov'

        subject = title
        from_addr = f"noreply@{agency_domain}"
        to_addr = f"recipient@{agency_domain}"

        # Build body text
        lines = []
        if has_cui and classification:
            lines.append(classification)
            lines.append('')
        lines.append(f"Subject: {title}")
        lines.append('')
        lines.append('Please find the attached document for your review.')
        lines.append('')
        lines.append(f"Organization: {agency}")
        if doc_data.get('document_date'):
            lines.append(f"Date: {doc_data['document_date']}")
        lines.append('')
        if has_cui:
            notice = doc_data.get('confidentiality_notice', '')
            if notice:
                lines.append('---')
                lines.append(f'CONFIDENTIALITY NOTICE: {notice}')

        plain_body = '\n'.join(lines)

        # Generate attachment (20% ZIP, 80% single doc)
        use_zip = random.random() < 0.2

        if use_zip:
            att_data = self._create_cui_zip_in_memory(doc_data, is_positive)
            att_filename = f"CUI_Documents_{doc_data.get('document_id', 'pkg')}.zip"
            attachments = [(att_data, att_filename, 'zip')]
        elif random.random() < 0.5:
            att_data = self._generate_cui_pdf_in_memory(doc_data)
            att_filename = f"{title.replace(' ', '_')[:30]}.pdf"
            attachments = [(att_data, att_filename, 'pdf')]
        else:
            att_data = self._generate_cui_docx_in_memory(doc_data)
            att_filename = f"{title.replace(' ', '_')[:30]}.docx"
            attachments = [(att_data, att_filename,
                           'vnd.openxmlformats-officedocument.wordprocessingml.document')]

        return self._build_and_save_email(
            subject=subject,
            from_addr=from_addr,
            to_addr=to_addr,
            plain_body=plain_body,
            attachments=attachments,
            filename=filename,
            message_id_domain=agency_domain,
        )

    def _generate_cui_pdf_in_memory(self, doc_data: Dict[str, Any]) -> bytes:
        """Generate CUI PDF document in memory using ReportLab."""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        # Classification header
        classification = doc_data.get('classification', '')
        if doc_data.get('has_cui', False) and classification:
            story.append(Paragraph(
                f'<para alignment="center"><font color="darkred"><b>{classification}</b></font></para>',
                styles['Normal'],
            ))
            story.append(Spacer(1, 0.2 * inch))

        # Title
        story.append(Paragraph(
            f"<b>{doc_data.get('title', 'Document')}</b>",
            styles['Title'],
        ))
        story.append(Spacer(1, 0.2 * inch))

        # Metadata table
        meta_data = [
            ['Document Information', ''],
            ['Agency:', doc_data.get('agency', '')],
            ['Date:', doc_data.get('document_date', '')],
            ['Category:', doc_data.get('category', '').replace('_', ' ').title()],
            ['Type:', doc_data.get('document_type', '').replace('_', ' ').title()],
        ]
        meta_table = Table(meta_data, colWidths=[2 * inch, 4 * inch])
        meta_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        story.append(meta_table)
        story.append(Spacer(1, 0.3 * inch))

        # Content
        skip_fields = {'document_id', 'document_type', 'category', 'subcategory',
                       'has_cui', 'classification', 'generated_date', 'title',
                       'confidentiality_notice', 'document_date', 'agency', 'authority'}

        for key, value in doc_data.items():
            if key in skip_fields:
                continue
            if isinstance(value, (str, int, float)) and value:
                story.append(Paragraph(
                    f"<b>{key.replace('_', ' ').title()}:</b> {value}",
                    styles['Normal'],
                ))
                story.append(Spacer(1, 0.1 * inch))

        # Confidentiality notice
        if doc_data.get('has_cui', False):
            notice = doc_data.get('confidentiality_notice', '')
            if notice:
                story.append(Spacer(1, 0.3 * inch))
                story.append(Paragraph(
                    f'<font size="9"><i>CONFIDENTIALITY NOTICE: {notice}</i></font>',
                    styles['Normal'],
                ))

        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()

    def _generate_cui_docx_in_memory(self, doc_data: Dict[str, Any]) -> bytes:
        """Generate CUI DOCX document in memory."""
        buffer = io.BytesIO()
        doc = Document()

        # Classification header
        classification = doc_data.get('classification', '')
        if doc_data.get('has_cui', False) and classification:
            header = doc.add_paragraph()
            header.alignment = 1  # Center
            run = header.add_run(classification)
            run.bold = True
            from docx.shared import RGBColor as DocxRGB
            run.font.color.rgb = DocxRGB(139, 0, 0)

        # Title
        doc.add_heading(doc_data.get('title', 'Document'), 0)

        # Metadata
        meta_fields = [
            ('Agency', doc_data.get('agency', '')),
            ('Date', doc_data.get('document_date', '')),
            ('Category', doc_data.get('category', '').replace('_', ' ').title()),
        ]
        for label, value in meta_fields:
            if value:
                doc.add_paragraph(f"{label}: {value}")

        doc.add_paragraph()

        # Content
        skip_fields = {'document_id', 'document_type', 'category', 'subcategory',
                       'has_cui', 'classification', 'generated_date', 'title',
                       'confidentiality_notice', 'document_date', 'agency', 'authority'}

        for key, value in doc_data.items():
            if key in skip_fields:
                continue
            if isinstance(value, (str, int, float)) and value:
                doc.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")

        # Notice
        if doc_data.get('has_cui', False):
            notice = doc_data.get('confidentiality_notice', '')
            if notice:
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.add_run('CONFIDENTIALITY NOTICE: ').bold = True
                p.add_run(notice)

        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _create_cui_zip_in_memory(self, doc_data: Dict[str, Any], is_positive: bool) -> bytes:
        """Create ZIP with 2-3 CUI documents in memory."""
        buffer = io.BytesIO()

        with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
            pdf_data = self._generate_cui_pdf_in_memory(doc_data)
            zipf.writestr("CUI_Document.pdf", pdf_data)

            docx_data = self._generate_cui_docx_in_memory(doc_data)
            zipf.writestr("CUI_Document.docx", docx_data)

            if random.random() < 0.5:
                zipf.writestr("CUI_Document_Copy.pdf", pdf_data)

        buffer.seek(0)
        return buffer.getvalue()
