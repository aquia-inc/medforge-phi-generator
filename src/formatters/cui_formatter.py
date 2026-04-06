"""
CUI document formatter.

Creates CUI documents in various formats (DOCX, PDF, EML, XLSX)
with proper classification headers, footers, and markings.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from formatters.base_email_formatter import BaseEmailFormatter
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment
from reportlab.lib.enums import TA_LEFT, TA_CENTER
import os
from typing import Any, Dict, Optional

from templates.components import ComponentConfiguration, get_docx_font_name, get_pdf_font_name

# Alignment map for component configs -> python-docx
_DOCX_ALIGNMENT_MAP = {
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "split": WD_ALIGN_PARAGRAPH.LEFT,
}

# Alignment map for component configs -> reportlab
_PDF_ALIGNMENT_MAP = {
    "center": TA_CENTER,
    "left": TA_LEFT,
    "split": TA_LEFT,
}


class CUIDocxFormatter:
    """Creates DOCX documents with CUI markings and content."""

    def __init__(self, output_dir: str = 'output'):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def create_cui_document(self, doc_data: Dict[str, Any], filename: str,
                            component_config: Optional[ComponentConfiguration] = None) -> str:
        """
        Create a DOCX document from CUI data with optional visual variation.

        Args:
            doc_data: Dictionary containing CUI document data
            filename: Output filename
            component_config: Optional component configuration for visual variety

        Returns:
            Path to created file
        """
        doc = Document()

        # Extract style settings from component config
        style_cfg = component_config.style.get_config() if component_config else None
        header_cfg = component_config.header.get_config() if component_config else None

        # Add classification header if CUI positive
        if doc_data.get('has_cui', False):
            self._add_classification_header(doc, doc_data, header_cfg=header_cfg)

        # Add document title
        self._add_title(doc, doc_data, style_cfg=style_cfg, header_cfg=header_cfg)

        # Add metadata section
        self._add_metadata_section(doc, doc_data, style_cfg=style_cfg)

        # Add main content based on document type
        self._add_content(doc, doc_data, style_cfg=style_cfg)

        # Add confidentiality notice if CUI positive
        if doc_data.get('has_cui', False):
            self._add_confidentiality_notice(doc, doc_data, style_cfg=style_cfg)

        # Add classification footer
        if doc_data.get('has_cui', False):
            self._add_classification_footer(doc, doc_data)

        # Save document
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        return filepath

    def _add_classification_header(self, doc: Document, doc_data: Dict[str, Any],
                                    header_cfg: Optional[Dict] = None):
        """Add classification banner at top of document."""
        classification = doc_data.get('classification', '')
        if not classification:
            return

        alignment = WD_ALIGN_PARAGRAPH.CENTER
        if header_cfg:
            alignment = _DOCX_ALIGNMENT_MAP.get(header_cfg.get("alignment", "center"),
                                                 WD_ALIGN_PARAGRAPH.CENTER)

        header = doc.add_paragraph()
        header.alignment = alignment
        run = header.add_run(classification)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(139, 0, 0)
        doc.add_paragraph()

    def _add_title(self, doc: Document, doc_data: Dict[str, Any],
                    style_cfg: Optional[Dict] = None, header_cfg: Optional[Dict] = None):
        """Add document title with optional style variation."""
        alignment = WD_ALIGN_PARAGRAPH.CENTER
        font_name = None
        font_size = Pt(14)
        if style_cfg:
            font_name = get_docx_font_name(style_cfg.get("font_family", "Arial"))
            font_size = Pt(style_cfg.get("font_size_title", 14))
        if header_cfg:
            alignment = _DOCX_ALIGNMENT_MAP.get(header_cfg.get("alignment", "center"),
                                                 WD_ALIGN_PARAGRAPH.CENTER)

        title = doc.add_paragraph()
        title.alignment = alignment
        run = title.add_run(doc_data.get('title', 'Document'))
        run.bold = True
        run.font.size = font_size
        if font_name:
            run.font.name = font_name
        if style_cfg and style_cfg.get("use_colors"):
            try:
                color = style_cfg.get("color_primary", "#000000").lstrip("#")
                run.font.color.rgb = RGBColor(int(color[:2], 16), int(color[2:4], 16), int(color[4:6], 16))
            except (ValueError, IndexError):
                pass
        doc.add_paragraph()

    def _add_metadata_section(self, doc: Document, doc_data: Dict[str, Any],
                               style_cfg: Optional[Dict] = None):
        """Add document metadata section with optional style variation."""
        metadata_items = []

        if 'agency' in doc_data:
            metadata_items.append(('Organization:', doc_data['agency']))
        if 'document_date' in doc_data:
            metadata_items.append(('Date:', doc_data['document_date']))
        if 'document_id' in doc_data:
            metadata_items.append(('Document ID:', doc_data['document_id']))
        if doc_data.get('authority'):
            metadata_items.append(('Authority:', doc_data['authority']))

        if not metadata_items:
            return

        font_name = get_docx_font_name(style_cfg["font_family"]) if style_cfg else None
        font_size = Pt(style_cfg.get("font_size_body", 11)) if style_cfg else None

        table = doc.add_table(rows=len(metadata_items), cols=2)
        for i, (label, value) in enumerate(metadata_items):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(value)
            label_run = table.rows[i].cells[0].paragraphs[0].runs[0]
            label_run.bold = True
            if font_name:
                label_run.font.name = font_name
                for run in table.rows[i].cells[1].paragraphs[0].runs:
                    run.font.name = font_name
            if font_size:
                label_run.font.size = font_size
                for run in table.rows[i].cells[1].paragraphs[0].runs:
                    run.font.size = font_size
        doc.add_paragraph()

    def _add_content(self, doc: Document, doc_data: Dict[str, Any],
                      style_cfg: Optional[Dict] = None):
        """Add main document content based on document type."""
        doc_type = doc_data.get('document_type', '')

        # Handle different document types
        if doc_type == 'coop_plan':
            self._add_coop_plan_content(doc, doc_data, style_cfg=style_cfg)
        elif doc_type == 'vulnerability_alert':
            self._add_vulnerability_content(doc, doc_data, style_cfg=style_cfg)
        elif doc_type == 'budget_memo':
            self._add_budget_memo_content(doc, doc_data, style_cfg=style_cfg)
        elif doc_type == 'attorney_memo':
            self._add_attorney_memo_content(doc, doc_data, style_cfg=style_cfg)
        elif doc_type in ['source_selection_plan', 'evaluation_report', 'igce']:
            self._add_procurement_content(doc, doc_data, style_cfg=style_cfg)
        elif doc_type == 'taxpayer_record':
            self._add_tax_record_content(doc, doc_data, style_cfg=style_cfg)
        elif doc_type == 'sar':
            self._add_sar_content(doc, doc_data, style_cfg=style_cfg)
        elif doc_type == 'written_determination':
            self._add_written_determination_content(doc, doc_data, style_cfg=style_cfg)
        elif doc_type == 'comptroller_report':
            self._add_comptroller_content(doc, doc_data, style_cfg=style_cfg)
        elif doc_type == 'bargaining_proposal':
            self._add_bargaining_content(doc, doc_data, style_cfg=style_cfg)
        elif doc_type == 'congressional_testimony':
            self._add_testimony_content(doc, doc_data, style_cfg=style_cfg)
        elif doc_type == 'retirement_estimate':
            self._add_retirement_content(doc, doc_data, style_cfg=style_cfg)
        else:
            self._add_generic_content(doc, doc_data, style_cfg=style_cfg)

    def _add_narrative_paragraphs(self, doc: Document, text: str,
                                   style_cfg: Optional[Dict] = None):
        """Split multi-paragraph text and render each as a styled paragraph."""
        if not text:
            return
        font_name = get_docx_font_name(style_cfg["font_family"]) if style_cfg else None
        font_size = Pt(style_cfg.get("font_size_body", 11)) if style_cfg else None
        line_spacing = style_cfg.get("line_height", 1.15) if style_cfg else None

        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        if not paragraphs:
            paragraphs = [text]

        for para_text in paragraphs:
            p = doc.add_paragraph()
            run = p.add_run(para_text)
            if font_name:
                run.font.name = font_name
            if font_size:
                run.font.size = font_size
            if line_spacing:
                p.paragraph_format.line_spacing = line_spacing

    def _add_styled_heading(self, doc: Document, text: str,
                             style_cfg: Optional[Dict] = None):
        """Add a section heading with optional style variation."""
        p = doc.add_paragraph(text, style='Heading 2')
        if style_cfg:
            for run in p.runs:
                font_name = get_docx_font_name(style_cfg.get("font_family", "Arial"))
                run.font.name = font_name
                if style_cfg.get("use_colors"):
                    try:
                        color = style_cfg.get("color_primary", "#000000").lstrip("#")
                        run.font.color.rgb = RGBColor(
                            int(color[:2], 16), int(color[2:4], 16), int(color[4:6], 16))
                    except (ValueError, IndexError):
                        pass

    def _add_coop_plan_content(self, doc: Document, doc_data: Dict[str, Any],
                                style_cfg: Optional[Dict] = None):
        """Add COOP plan specific content."""
        self._add_styled_heading(doc, 'EXECUTIVE SUMMARY', style_cfg=style_cfg)
        self._add_narrative_paragraphs(doc, doc_data.get('executive_summary', ''),
                                       style_cfg=style_cfg)

        # Essential Functions
        if 'essential_functions' in doc_data:
            doc.add_paragraph('ESSENTIAL FUNCTIONS PRIORITIZATION', style='Heading 2')
            for func in doc_data['essential_functions']:
                doc.add_paragraph(f"• Priority {func['priority']}: {func['function']}")

        # Alternate Locations
        if 'alternate_locations' in doc_data:
            doc.add_paragraph('ALTERNATE LOCATIONS', style='Heading 2')
            locs = doc_data['alternate_locations']
            doc.add_paragraph(f"Primary: {locs.get('primary', 'N/A')}")
            doc.add_paragraph(f"Secondary: {locs.get('secondary', 'N/A')}")
            doc.add_paragraph(f"Devolution: {locs.get('devolution_distance', 'N/A')}")

        # Activation Triggers
        if 'activation_triggers' in doc_data:
            doc.add_paragraph('ACTIVATION TRIGGERS', style='Heading 2')
            for trigger in doc_data['activation_triggers']:
                doc.add_paragraph(f"• {trigger}")

        # ERG Details
        if 'erg_details' in doc_data:
            doc.add_paragraph('EMERGENCY RELOCATION GROUP', style='Heading 2')
            erg = doc_data['erg_details']
            doc.add_paragraph(f"ERG Leader: {erg.get('leader_name', '')} ({erg.get('leader_title', '')})")
            doc.add_paragraph(f"ERG Size: {erg.get('size', '')} personnel")
            doc.add_paragraph(f"Deployment Time: Within {erg.get('deployment_hours', '')} hours")

    def _add_vulnerability_content(self, doc: Document, doc_data: Dict[str, Any],
                                     style_cfg: Optional[Dict] = None):
        """Add vulnerability alert content."""
        self._add_styled_heading(doc, 'ALERT DETAILS', style_cfg=style_cfg)
        doc.add_paragraph(f"Alert ID: {doc_data.get('alert_id', '')}")
        doc.add_paragraph(f"Severity: {doc_data.get('severity', '')}")
        doc.add_paragraph(f"CVSS Score: {doc_data.get('cvss_score', '')}")
        doc.add_paragraph(f"CVE: {doc_data.get('cve_id', '')}")

        self._add_styled_heading(doc, 'AFFECTED SYSTEM', style_cfg=style_cfg)
        doc.add_paragraph(f"System: {doc_data.get('affected_system', '')}")
        doc.add_paragraph(f"Versions: {doc_data.get('affected_versions', '')}")

        self._add_styled_heading(doc, 'DESCRIPTION', style_cfg=style_cfg)
        self._add_narrative_paragraphs(doc, doc_data.get('description', ''),
                                       style_cfg=style_cfg)

        if 'remediation' in doc_data:
            self._add_styled_heading(doc, 'REMEDIATION', style_cfg=style_cfg)
            rem = doc_data['remediation']
            self._add_narrative_paragraphs(doc, rem.get('action', ''), style_cfg=style_cfg)
            doc.add_paragraph(f"Target Version: {rem.get('target_version', '')}")
            doc.add_paragraph(f"Deadline: {rem.get('deadline', '')}")

    def _add_budget_memo_content(self, doc: Document, doc_data: Dict[str, Any],
                                  style_cfg: Optional[Dict] = None):
        """Add budget memo content."""
        doc.add_paragraph(f"TO: {doc_data.get('to', '')}")
        doc.add_paragraph(f"FROM: {doc_data.get('from', '')}")
        doc.add_paragraph(f"SUBJECT: {doc_data.get('subject', '')}")
        doc.add_paragraph()

        self._add_styled_heading(doc, 'PRESIDENTIAL DECISION:', style_cfg=style_cfg)
        self._add_narrative_paragraphs(doc, doc_data.get('decision', ''),
                                       style_cfg=style_cfg)

        if 'key_decision_points' in doc_data:
            self._add_styled_heading(doc, 'KEY DECISION POINTS:', style_cfg=style_cfg)
            for point in doc_data['key_decision_points']:
                doc.add_paragraph(f"\u2022 {point}")

    def _add_attorney_memo_content(self, doc: Document, doc_data: Dict[str, Any],
                                     style_cfg: Optional[Dict] = None):
        """Add attorney memorandum content."""
        if 'privilege_assertion' in doc_data:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(' | '.join(doc_data['privilege_assertion']))
            run.italic = True

        attorney = doc_data.get('attorney', {})
        client = doc_data.get('client', {})
        doc.add_paragraph(f"TO: {client.get('name', '')} - {client.get('title', '')}")
        doc.add_paragraph(f"FROM: {attorney.get('name', '')} - {attorney.get('title', '')}")
        doc.add_paragraph(f"RE: {doc_data.get('subject', '')}")
        doc.add_paragraph()

        self._add_styled_heading(doc, 'QUESTION PRESENTED:', style_cfg=style_cfg)
        self._add_narrative_paragraphs(doc, doc_data.get('question_presented', ''),
                                       style_cfg=style_cfg)

        self._add_styled_heading(doc, 'BRIEF ANSWER:', style_cfg=style_cfg)
        self._add_narrative_paragraphs(doc, doc_data.get('brief_answer', ''),
                                       style_cfg=style_cfg)

        self._add_styled_heading(doc, 'ANALYSIS:', style_cfg=style_cfg)
        self._add_narrative_paragraphs(doc, doc_data.get('analysis', ''),
                                       style_cfg=style_cfg)

        self._add_styled_heading(doc, 'RECOMMENDATION:', style_cfg=style_cfg)
        self._add_narrative_paragraphs(doc, doc_data.get('recommendation', ''),
                                       style_cfg=style_cfg)

    def _add_procurement_content(self, doc: Document, doc_data: Dict[str, Any],
                                  style_cfg: Optional[Dict] = None):
        """Add procurement document content."""
        doc.add_paragraph(f"Solicitation Number: {doc_data.get('solicitation_number', '')}")

        if doc_data.get('document_type') == 'source_selection_plan':
            doc.add_paragraph(f"Program: {doc_data.get('program', '')}")
            doc.add_paragraph(f"Estimated Value: {doc_data.get('estimated_value', '')}")
            doc.add_paragraph(f"Contract Type: {doc_data.get('contract_type', '')}")

            if 'evaluation_factors' in doc_data:
                self._add_styled_heading(doc, 'EVALUATION FACTORS:', style_cfg=style_cfg)
                for factor in doc_data['evaluation_factors']:
                    doc.add_paragraph(f"\u2022 {factor['factor']}: Weight {factor['weight']}%")

            if doc_data.get('executive_summary'):
                self._add_styled_heading(doc, 'ACQUISITION SUMMARY:', style_cfg=style_cfg)
                self._add_narrative_paragraphs(doc, doc_data['executive_summary'],
                                               style_cfg=style_cfg)
            if doc_data.get('justification'):
                self._add_styled_heading(doc, 'JUSTIFICATION:', style_cfg=style_cfg)
                self._add_narrative_paragraphs(doc, doc_data['justification'],
                                               style_cfg=style_cfg)

        elif doc_data.get('document_type') == 'evaluation_report':
            doc.add_paragraph(f"Offeror: {doc_data.get('offeror', '')}")
            doc.add_paragraph(f"Overall Rating: {doc_data.get('overall_rating', '')}")

    def _add_tax_record_content(self, doc: Document, doc_data: Dict[str, Any],
                                 style_cfg: Optional[Dict] = None):
        """Add tax record content."""
        taxpayer = doc_data.get('taxpayer', {})
        doc.add_paragraph(f"Taxpayer: {taxpayer.get('name', '')}")
        doc.add_paragraph(f"SSN: {taxpayer.get('ssn_masked', '')}")
        doc.add_paragraph(f"Tax Period: {doc_data.get('tax_period', '')}")

        if 'account_summary' in doc_data:
            doc.add_paragraph('ACCOUNT SUMMARY:', style='Heading 2')
            summary = doc_data['account_summary']
            for key, value in summary.items():
                doc.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")

    def _add_sar_content(self, doc: Document, doc_data: Dict[str, Any],
                          style_cfg: Optional[Dict] = None):
        """Add Suspicious Activity Report content."""
        doc.add_paragraph(f"SAR Number: {doc_data.get('sar_number', '')}")
        doc.add_paragraph(f"Filing Date: {doc_data.get('filing_date', '')}")

        inst = doc_data.get('filing_institution', {})
        if isinstance(inst, dict):
            self._add_styled_heading(doc, 'FILING INSTITUTION', style_cfg=style_cfg)
            doc.add_paragraph(f"Name: {inst.get('name', '')}")
            doc.add_paragraph(f"RSSD ID: {inst.get('rssd_id', '')}")

        subject = doc_data.get('subject', {})
        if isinstance(subject, dict):
            self._add_styled_heading(doc, 'SUBJECT INFORMATION', style_cfg=style_cfg)
            doc.add_paragraph(f"Name: {subject.get('name', '')}")
            doc.add_paragraph(f"DOB: {subject.get('dob', '')}")
            doc.add_paragraph(f"SSN (last 4): {subject.get('ssn_last4', '')}")

        activity = doc_data.get('suspicious_activity', {})
        if isinstance(activity, dict):
            self._add_styled_heading(doc, 'SUSPICIOUS ACTIVITY', style_cfg=style_cfg)
            doc.add_paragraph(f"Type: {activity.get('type', '')}")
            doc.add_paragraph(f"Amount: {activity.get('amount_formatted', activity.get('amount', ''))}")
            date_range = activity.get('date_range', {})
            if isinstance(date_range, dict):
                doc.add_paragraph(f"Period: {date_range.get('start', '')} - {date_range.get('end', '')}")
            narrative = activity.get('narrative', '')
            if narrative:
                self._add_styled_heading(doc, 'NARRATIVE', style_cfg=style_cfg)
                self._add_narrative_paragraphs(doc, narrative, style_cfg=style_cfg)

    def _add_written_determination_content(self, doc: Document, doc_data: Dict[str, Any],
                                            style_cfg: Optional[Dict] = None):
        """Add IRS written determination content (IRAC format)."""
        doc.add_paragraph(f"Document Number: {doc_data.get('document_number', '')}")
        doc.add_paragraph(f"Issue Date: {doc_data.get('issue_date', '')}")

        det_type = doc_data.get('determination_type', '')
        if det_type:
            doc.add_paragraph(f"Type: {det_type}")

        issues = doc_data.get('issues', [])
        if issues:
            self._add_styled_heading(doc, 'ISSUES', style_cfg=style_cfg)
            for issue in (issues if isinstance(issues, list) else [issues]):
                doc.add_paragraph(f"\u2022 {issue}")

        code_sections = doc_data.get('code_sections', [])
        if code_sections:
            self._add_styled_heading(doc, 'APPLICABLE CODE SECTIONS', style_cfg=style_cfg)
            for sec in (code_sections if isinstance(code_sections, list) else [code_sections]):
                doc.add_paragraph(f"\u2022 {sec}")

        self._add_styled_heading(doc, 'FACTS', style_cfg=style_cfg)
        self._add_narrative_paragraphs(doc, doc_data.get('facts', ''), style_cfg=style_cfg)

        self._add_styled_heading(doc, 'LAW AND ANALYSIS', style_cfg=style_cfg)
        self._add_narrative_paragraphs(doc, doc_data.get('law_and_analysis', ''), style_cfg=style_cfg)

        self._add_styled_heading(doc, 'CONCLUSION', style_cfg=style_cfg)
        self._add_narrative_paragraphs(doc, doc_data.get('conclusion', ''), style_cfg=style_cfg)

        caveats = doc_data.get('caveats', '')
        if caveats:
            self._add_styled_heading(doc, 'CAVEATS', style_cfg=style_cfg)
            self._add_narrative_paragraphs(doc, caveats, style_cfg=style_cfg)

    def _add_comptroller_content(self, doc: Document, doc_data: Dict[str, Any],
                                  style_cfg: Optional[Dict] = None):
        """Add GAO comptroller report content."""
        doc.add_paragraph(f"Report Number: {doc_data.get('report_number', '')}")
        doc.add_paragraph(f"Agency Reviewed: {doc_data.get('agency_reviewed', '')}")
        doc.add_paragraph(f"Topic: {doc_data.get('topic', '')}")
        doc.add_paragraph(f"Status: {doc_data.get('report_status', '')}")

        review = doc_data.get('review_period', {})
        if isinstance(review, dict):
            doc.add_paragraph(f"Review Period: {review.get('start', '')} - {review.get('end', '')}")

        if doc_data.get('executive_summary'):
            self._add_styled_heading(doc, 'EXECUTIVE SUMMARY', style_cfg=style_cfg)
            self._add_narrative_paragraphs(doc, doc_data['executive_summary'], style_cfg=style_cfg)

        findings = doc_data.get('findings', [])
        if findings:
            self._add_styled_heading(doc, 'FINDINGS', style_cfg=style_cfg)
            for f in findings:
                if isinstance(f, dict):
                    doc.add_paragraph(f"Finding {f.get('finding_number', '')}: {f.get('description', '')}")
                    if f.get('significance'):
                        doc.add_paragraph(f"  Significance: {f['significance']}")
                else:
                    doc.add_paragraph(f"\u2022 {f}")

        doc.add_paragraph(f"Recommendations: {doc_data.get('recommendations', '')}")
        doc.add_paragraph(f"Estimated Savings: {doc_data.get('estimated_savings', '')}")

    def _add_bargaining_content(self, doc: Document, doc_data: Dict[str, Any],
                                 style_cfg: Optional[Dict] = None):
        """Add collective bargaining proposal content."""
        doc.add_paragraph(f"Proposal Number: {doc_data.get('proposal_number', '')}")
        doc.add_paragraph(f"Union: {doc_data.get('union', '')}")
        doc.add_paragraph(f"Negotiation Status: {doc_data.get('negotiation_status', '')}")
        doc.add_paragraph(f"Date Submitted: {doc_data.get('date_submitted', '')}")

        articles = doc_data.get('articles_under_negotiation', [])
        if articles:
            self._add_styled_heading(doc, 'ARTICLES UNDER NEGOTIATION', style_cfg=style_cfg)
            for art in articles:
                if isinstance(art, dict):
                    doc.add_paragraph(f"\u2022 {art.get('article', '')}: {art.get('subject', '')}")
                else:
                    doc.add_paragraph(f"\u2022 {art}")

        rules = doc_data.get('ground_rules', {})
        if isinstance(rules, dict):
            self._add_styled_heading(doc, 'GROUND RULES', style_cfg=style_cfg)
            doc.add_paragraph(f"Location: {rules.get('negotiation_location', '')}")
            dates = rules.get('session_dates', [])
            if dates:
                doc.add_paragraph(f"Sessions: {', '.join(str(d) for d in dates)}")

        for role in ('management_team_lead', 'union_team_lead'):
            lead = doc_data.get(role, {})
            if isinstance(lead, dict) and lead.get('name'):
                label = role.replace('_', ' ').title()
                doc.add_paragraph(f"{label}: {lead['name']} - {lead.get('title', '')}")

    def _add_testimony_content(self, doc: Document, doc_data: Dict[str, Any],
                                style_cfg: Optional[Dict] = None):
        """Add congressional testimony content."""
        status = doc_data.get('document_status', '')
        if status:
            doc.add_paragraph(f"Status: {status}")

        witness = doc_data.get('witness', {})
        if isinstance(witness, dict):
            self._add_styled_heading(doc, 'WITNESS', style_cfg=style_cfg)
            doc.add_paragraph(f"Name: {witness.get('name', '')}")
            doc.add_paragraph(f"Title: {witness.get('title', '')}")
            doc.add_paragraph(f"Agency: {witness.get('agency', '')}")

        doc.add_paragraph(f"Committee: {doc_data.get('committee', '')}")
        doc.add_paragraph(f"Hearing Topic: {doc_data.get('hearing_topic', '')}")
        doc.add_paragraph(f"Scheduled Date: {doc_data.get('scheduled_date', '')}")

        summary = doc_data.get('prepared_statement_summary', '')
        if summary:
            self._add_styled_heading(doc, 'PREPARED STATEMENT SUMMARY', style_cfg=style_cfg)
            self._add_narrative_paragraphs(doc, summary, style_cfg=style_cfg)

        messages = doc_data.get('key_messages', [])
        if messages:
            self._add_styled_heading(doc, 'KEY MESSAGES', style_cfg=style_cfg)
            for msg in messages:
                doc.add_paragraph(f"\u2022 {msg}")

        doc.add_paragraph(f"Clearance Status: {doc_data.get('clearance_status', '')}")

    def _add_retirement_content(self, doc: Document, doc_data: Dict[str, Any],
                                 style_cfg: Optional[Dict] = None):
        """Add retirement estimate content."""
        doc.add_paragraph(f"Estimate Number: {doc_data.get('estimate_number', '')}")

        emp = doc_data.get('employee', {})
        if isinstance(emp, dict):
            self._add_styled_heading(doc, 'EMPLOYEE INFORMATION', style_cfg=style_cfg)
            doc.add_paragraph(f"Name: {emp.get('name', '')}")
            doc.add_paragraph(f"Employee ID: {emp.get('employee_id', '')}")
            doc.add_paragraph(f"Agency: {emp.get('agency', '')}")
            doc.add_paragraph(f"Grade: {emp.get('grade', '')}")

        svc = doc_data.get('service_computation', {})
        if isinstance(svc, dict):
            self._add_styled_heading(doc, 'SERVICE COMPUTATION', style_cfg=style_cfg)
            doc.add_paragraph(f"Retirement System: {svc.get('retirement_system', '')}")
            doc.add_paragraph(f"Years of Service: {svc.get('years_of_service', '')}")
            doc.add_paragraph(f"Sick Leave Credit: {svc.get('sick_leave_credit_months', '')} months")

        salary = doc_data.get('salary_information', {})
        if isinstance(salary, dict):
            self._add_styled_heading(doc, 'SALARY INFORMATION', style_cfg=style_cfg)
            doc.add_paragraph(f"Current Salary: {salary.get('current_salary', '')}")
            doc.add_paragraph(f"High-3 Average: {salary.get('high_3_average', '')}")

        benefits = doc_data.get('estimated_benefits', {})
        if isinstance(benefits, dict):
            self._add_styled_heading(doc, 'ESTIMATED BENEFITS', style_cfg=style_cfg)
            doc.add_paragraph(f"Gross Monthly Annuity: {benefits.get('gross_monthly_annuity', '')}")
            if benefits.get('fers_supplement'):
                doc.add_paragraph(f"FERS Supplement: {benefits['fers_supplement']}")
            if benefits.get('tsp_balance'):
                doc.add_paragraph(f"TSP Balance: {benefits['tsp_balance']}")

        if doc_data.get('projected_retirement_date'):
            doc.add_paragraph(f"Projected Retirement Date: {doc_data['projected_retirement_date']}")

        if doc_data.get('executive_summary'):
            self._add_styled_heading(doc, 'SUMMARY', style_cfg=style_cfg)
            self._add_narrative_paragraphs(doc, doc_data['executive_summary'], style_cfg=style_cfg)

    def _add_generic_content(self, doc: Document, doc_data: Dict[str, Any],
                              style_cfg: Optional[Dict] = None):
        """Add generic document content with optional narrative rendering."""
        skip_fields = {'document_id', 'document_type', 'category', 'subcategory',
                       'has_cui', 'classification', 'authority', 'distribution',
                       'generated_date', 'document_date', 'agency', 'title',
                       'confidentiality_notice'}
        # Fields that contain LLM narrative text (render as multi-paragraph)
        narrative_fields = {'executive_summary', 'body_content', 'analysis',
                           'recommendations', 'risk_assessment', 'justification',
                           'description', 'body'}

        for key, value in doc_data.items():
            if key in skip_fields:
                continue

            if isinstance(value, str):
                if key in narrative_fields and '\n' in value:
                    self._add_styled_heading(doc, key.replace('_', ' ').title(),
                                             style_cfg=style_cfg)
                    self._add_narrative_paragraphs(doc, value, style_cfg=style_cfg)
                else:
                    doc.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")
            elif isinstance(value, list):
                self._add_styled_heading(doc, key.replace('_', ' ').title(),
                                         style_cfg=style_cfg)
                for item in value:
                    doc.add_paragraph(f"\u2022 {item}")
            elif isinstance(value, dict):
                self._add_styled_heading(doc, key.replace('_', ' ').title(),
                                         style_cfg=style_cfg)
                for k, v in value.items():
                    doc.add_paragraph(f"  {k}: {v}")

    def _add_confidentiality_notice(self, doc: Document, doc_data: Dict[str, Any],
                                     style_cfg: Optional[Dict] = None):
        """Add confidentiality notice at bottom of document."""
        notice_text = doc_data.get('confidentiality_notice', '')
        if not notice_text:
            return

        font_name = get_docx_font_name(style_cfg["font_family"]) if style_cfg else None

        doc.add_paragraph()
        notice = doc.add_paragraph()
        notice.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = notice.add_run('CONFIDENTIALITY NOTICE:')
        run.bold = True
        run.font.size = Pt(10)
        if font_name:
            run.font.name = font_name

        notice_para = doc.add_paragraph()
        nr = notice_para.add_run(notice_text)
        nr.font.size = Pt(9)
        nr.italic = True
        if font_name:
            nr.font.name = font_name

    def _add_classification_footer(self, doc: Document, doc_data: Dict[str, Any]):
        """Add classification footer."""
        # Only add footer if classification is present and not empty
        classification = doc_data.get('classification', '')
        if not classification:
            return

        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run(classification)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(139, 0, 0)

        # Distribution statement
        if doc_data.get('distribution'):
            dist = doc.add_paragraph()
            dist.alignment = WD_ALIGN_PARAGRAPH.CENTER
            dist_run = dist.add_run(doc_data['distribution'])
            dist_run.font.size = Pt(9)
            dist_run.italic = True

    def _get_default_notice(self) -> str:
        return ("This document contains Controlled Unclassified Information (CUI) that requires "
                "safeguarding or dissemination controls pursuant to applicable laws, regulations, "
                "and Government-wide policies.")


class CUIEmailFormatter(BaseEmailFormatter):
    """Creates EML email files with CUI content."""

    def create_cui_email(self, doc_data: Dict[str, Any], filename: str) -> str:
        """
        Create an EML email from CUI data.

        Args:
            doc_data: Dictionary containing CUI document data
            filename: Output filename

        Returns:
            Path to created file
        """
        subject = doc_data.get('title', 'Document')

        agency = doc_data.get('agency', 'Department of Health and Human Services')
        agency_domain = agency.lower().replace(' ', '').replace('of', '')[:10] + '.gov'

        plain_text = self._build_plain_text(doc_data)
        html_text = self._build_html(doc_data)

        return self._build_and_save_email(
            subject=subject,
            from_addr=f"noreply@{agency_domain}",
            to_addr=f"recipient@{agency_domain}",
            plain_body=plain_text,
            html_body=html_text,
            filename=filename,
            message_id_domain=agency_domain,
        )

    def _build_plain_text(self, doc_data: Dict[str, Any]) -> str:
        """Build plain text email body."""
        lines = []

        # Only add classification if present and not empty
        classification = doc_data.get('classification', '')
        if doc_data.get('has_cui', False) and classification:
            lines.append(classification)
            lines.append('')

        lines.append(doc_data.get('title', 'Document'))
        lines.append('=' * 50)
        lines.append('')

        # Add document content based on type
        doc_type = doc_data.get('document_type', '')

        if doc_type == 'vulnerability_alert':
            lines.extend(self._format_vulnerability_alert_text(doc_data))
        elif doc_type == 'servicenow_ticket':
            lines.extend(self._format_servicenow_text(doc_data))
        else:
            lines.extend(self._format_generic_text(doc_data))

        if doc_data.get('has_cui', False):
            lines.append('')
            lines.append('-' * 50)
            lines.append('CONFIDENTIALITY NOTICE:')
            lines.append(doc_data.get('confidentiality_notice', ''))

        return '\n'.join(lines)

    def _format_vulnerability_alert_text(self, doc_data: Dict[str, Any]) -> list:
        """Format vulnerability alert for plain text."""
        lines = [
            f"Alert ID: {doc_data.get('alert_id', '')}",
            f"Severity: {doc_data.get('severity', '')}",
            f"CVSS Score: {doc_data.get('cvss_score', '')}",
            f"CVE: {doc_data.get('cve_id', '')}",
            '',
            f"Affected System: {doc_data.get('affected_system', '')}",
            f"Affected Versions: {doc_data.get('affected_versions', '')}",
            '',
            'Description:',
            doc_data.get('description', ''),
            '',
        ]
        if 'remediation' in doc_data:
            rem = doc_data['remediation']
            lines.extend([
                'Remediation:',
                f"  Action: {rem.get('action', '')}",
                f"  Target Version: {rem.get('target_version', '')}",
                f"  Deadline: {rem.get('deadline', '')}",
            ])
        return lines

    def _format_servicenow_text(self, doc_data: Dict[str, Any]) -> list:
        """Format ServiceNow ticket for plain text."""
        return [
            f"Ticket Number: {doc_data.get('ticket_number', '')}",
            f"Requester: {doc_data.get('requester', '')}",
            f"Department: {doc_data.get('department', '')}",
            f"Request Type: {doc_data.get('request_type', '')}",
            f"Priority: {doc_data.get('priority', '')}",
            f"Status: {doc_data.get('status', '')}",
            '',
            'Description:',
            doc_data.get('description', ''),
            '',
            f"Assigned To: {doc_data.get('assigned_to', '')}",
            f"Created: {doc_data.get('created_date', '')}",
        ]

    def _format_generic_text(self, doc_data: Dict[str, Any]) -> list:
        """Format generic document for plain text."""
        lines = []
        skip_fields = {'document_id', 'document_type', 'category', 'subcategory',
                       'has_cui', 'classification', 'generated_date', 'title',
                       'confidentiality_notice'}

        for key, value in doc_data.items():
            if key in skip_fields:
                continue
            if isinstance(value, (str, int, float)):
                lines.append(f"{key.replace('_', ' ').title()}: {value}")
        return lines

    def _build_html(self, doc_data: Dict[str, Any]) -> str:
        """Build HTML email body."""
        html_parts = ['<html><head></head><body style="font-family: Arial, sans-serif;">']

        # Only add classification if present and not empty
        classification = doc_data.get('classification', '')
        if doc_data.get('has_cui', False) and classification:
            html_parts.append(
                f'<p style="color: darkred; font-weight: bold; text-align: center;">'
                f'{classification}</p>'
            )

        html_parts.append(f'<h2>{doc_data.get("title", "Document")}</h2>')
        html_parts.append(f'<p><strong>Date:</strong> {doc_data.get("document_date", "")}</p>')
        html_parts.append(f'<p><strong>Organization:</strong> {doc_data.get("agency", "")}</p>')

        # Add content based on document type
        doc_type = doc_data.get('document_type', '')
        if doc_type == 'vulnerability_alert':
            html_parts.append(self._format_vulnerability_alert_html(doc_data))
        else:
            html_parts.append(self._format_generic_html(doc_data))

        if doc_data.get('has_cui', False):
            html_parts.append('<hr>')
            html_parts.append('<p style="font-size: 10px; font-style: italic;">')
            html_parts.append('<strong>CONFIDENTIALITY NOTICE:</strong><br>')
            html_parts.append(doc_data.get('confidentiality_notice', ''))
            html_parts.append('</p>')

        html_parts.append('</body></html>')
        return ''.join(html_parts)

    def _format_vulnerability_alert_html(self, doc_data: Dict[str, Any]) -> str:
        """Format vulnerability alert for HTML."""
        severity = doc_data.get('severity', 'Unknown')
        severity_color = {'Critical': 'red', 'High': 'orange', 'Medium': 'yellow', 'Low': 'green'}.get(severity, 'gray')

        return f'''
        <table border="1" cellpadding="5" style="border-collapse: collapse;">
            <tr><td><strong>Alert ID:</strong></td><td>{doc_data.get('alert_id', '')}</td></tr>
            <tr><td><strong>Severity:</strong></td><td style="background-color: {severity_color};">{severity}</td></tr>
            <tr><td><strong>CVSS Score:</strong></td><td>{doc_data.get('cvss_score', '')}</td></tr>
            <tr><td><strong>CVE:</strong></td><td>{doc_data.get('cve_id', '')}</td></tr>
            <tr><td><strong>Affected System:</strong></td><td>{doc_data.get('affected_system', '')}</td></tr>
        </table>
        <h3>Description</h3>
        <p>{doc_data.get('description', '')}</p>
        '''

    def _format_generic_html(self, doc_data: Dict[str, Any]) -> str:
        """Format generic document for HTML."""
        html = '<table border="1" cellpadding="5" style="border-collapse: collapse;">'
        skip_fields = {'document_id', 'document_type', 'category', 'subcategory',
                       'has_cui', 'classification', 'generated_date', 'title',
                       'confidentiality_notice', 'document_date', 'agency'}

        for key, value in doc_data.items():
            if key in skip_fields:
                continue
            if isinstance(value, (str, int, float)):
                html += f'<tr><td><strong>{key.replace("_", " ").title()}:</strong></td><td>{value}</td></tr>'

        html += '</table>'
        return html


class CUIPdfFormatter:
    """Creates PDF documents with CUI markings."""

    def __init__(self, output_dir: str = 'output'):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def create_cui_pdf(self, doc_data: Dict[str, Any], filename: str,
                        component_config: Optional[ComponentConfiguration] = None) -> str:
        """
        Create a PDF document from CUI data with optional visual variation.

        Args:
            doc_data: Dictionary containing CUI document data
            filename: Output filename
            component_config: Optional component configuration for visual variety

        Returns:
            Path to created file
        """
        filepath = os.path.join(self.output_dir, filename)
        doc = SimpleDocTemplate(filepath, pagesize=letter,
                               rightMargin=72, leftMargin=72,
                               topMargin=72, bottomMargin=72)

        styles = getSampleStyleSheet()
        story = []

        # Extract style settings from component config
        style_cfg = component_config.style.get_config() if component_config else None
        header_cfg = component_config.header.get_config() if component_config else None

        font_name = get_pdf_font_name(style_cfg.get("font_family", "Helvetica")) if style_cfg else "Helvetica"
        font_size = style_cfg.get("font_size_body", 10) if style_cfg else 10
        title_size = style_cfg.get("font_size_title", 16) if style_cfg else 16
        leading = font_size * (style_cfg.get("line_height", 1.4) if style_cfg else 1.4)
        pdf_alignment = TA_CENTER
        if header_cfg:
            pdf_alignment = _PDF_ALIGNMENT_MAP.get(
                header_cfg.get("alignment", "center"), TA_CENTER)

        body_style = ParagraphStyle(
            'CUIBody', parent=styles['Normal'],
            fontName=font_name, fontSize=font_size, leading=leading,
            spaceBefore=4, spaceAfter=6,
        )

        # Classification header
        classification = doc_data.get('classification', '')
        if doc_data.get('has_cui', False) and classification:
            cui_style = ParagraphStyle(
                'CUIHeader', parent=styles['Normal'],
                fontSize=12, textColor=colors.darkred,
                alignment=pdf_alignment, spaceAfter=20,
                fontName=f'{font_name}-Bold' if font_name in ('Helvetica', 'Times-Roman', 'Courier') else font_name,
            )
            story.append(Paragraph(classification, cui_style))

        # Title
        title_style = ParagraphStyle(
            'CUITitle', parent=styles['Heading1'],
            fontName=font_name, fontSize=title_size,
            alignment=pdf_alignment, spaceAfter=20,
        )
        story.append(Paragraph(doc_data.get('title', 'Document'), title_style))

        # Metadata
        story.append(Paragraph(f"<b>Organization:</b> {doc_data.get('agency', '')}", body_style))
        story.append(Paragraph(f"<b>Date:</b> {doc_data.get('document_date', '')}", body_style))
        if doc_data.get('authority'):
            story.append(Paragraph(f"<b>Authority:</b> {doc_data.get('authority', '')}", body_style))
        story.append(Spacer(1, 20))

        # Content
        self._add_pdf_content(story, doc_data, styles)

        # Confidentiality notice
        if doc_data.get('has_cui', False):
            notice_text = doc_data.get('confidentiality_notice', '')
            if notice_text:  # Only add if present
                story.append(Spacer(1, 30))
                notice_style = ParagraphStyle(
                    'Notice',
                    parent=styles['Normal'],
                    fontSize=9,
                    textColor=colors.gray,
                    spaceAfter=10
                )
                story.append(Paragraph('<b>CONFIDENTIALITY NOTICE:</b>', notice_style))
                story.append(Paragraph(notice_text, notice_style))

        doc.build(story)
        return filepath

    def _add_pdf_content(self, story: list, doc_data: Dict[str, Any], styles):
        """Add content to PDF story."""
        skip_fields = {'document_id', 'document_type', 'category', 'subcategory',
                       'has_cui', 'classification', 'authority', 'distribution',
                       'generated_date', 'document_date', 'agency', 'title',
                       'confidentiality_notice'}

        for key, value in doc_data.items():
            if key in skip_fields:
                continue

            if isinstance(value, str):
                story.append(Paragraph(f"<b>{key.replace('_', ' ').title()}:</b> {value}", styles['Normal']))
            elif isinstance(value, (int, float)):
                story.append(Paragraph(f"<b>{key.replace('_', ' ').title()}:</b> {value}", styles['Normal']))
            elif isinstance(value, list):
                story.append(Paragraph(f"<b>{key.replace('_', ' ').title()}:</b>", styles['Normal']))
                for item in value:
                    if isinstance(item, dict):
                        for k, v in item.items():
                            story.append(Paragraph(f"  • {k}: {v}", styles['Normal']))
                    else:
                        story.append(Paragraph(f"  • {item}", styles['Normal']))
            elif isinstance(value, dict):
                story.append(Paragraph(f"<b>{key.replace('_', ' ').title()}:</b>", styles['Normal']))
                for k, v in value.items():
                    story.append(Paragraph(f"  {k.replace('_', ' ').title()}: {v}", styles['Normal']))

            story.append(Spacer(1, 6))


class CUIXlsxFormatter:
    """Creates XLSX spreadsheets with CUI markings."""

    def __init__(self, output_dir: str = 'output'):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def create_cui_xlsx(self, doc_data: Dict[str, Any], filename: str) -> str:
        """
        Create an XLSX spreadsheet from CUI data.

        Args:
            doc_data: Dictionary containing CUI document data
            filename: Output filename

        Returns:
            Path to created file
        """
        wb = Workbook()
        ws = wb.active
        ws.title = "CUI Document"

        row = 1

        # Classification header - only add if classification is present and not empty
        classification = doc_data.get('classification', '')
        if doc_data.get('has_cui', False) and classification:
            ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
            cell = ws.cell(row=row, column=1, value=classification)
            cell.font = Font(bold=True, color='8B0000')
            cell.alignment = Alignment(horizontal='center')
            row += 2

        # Title
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
        cell = ws.cell(row=row, column=1, value=doc_data.get('title', 'Document'))
        cell.font = Font(bold=True, size=14)
        cell.alignment = Alignment(horizontal='center')
        row += 2

        # Metadata
        metadata = [
            ('Organization', doc_data.get('agency', '')),
            ('Date', doc_data.get('document_date', '')),
            ('Document ID', doc_data.get('document_id', '')),
        ]
        if doc_data.get('authority'):
            metadata.append(('Authority', doc_data.get('authority', '')))

        for label, value in metadata:
            ws.cell(row=row, column=1, value=label).font = Font(bold=True)
            ws.cell(row=row, column=2, value=value)
            row += 1

        row += 1

        # Content
        row = self._add_xlsx_content(ws, row, doc_data)

        # Confidentiality notice
        if doc_data.get('has_cui', False):
            notice_text = doc_data.get('confidentiality_notice', '')
            if notice_text:  # Only add if present
                row += 2
                ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
                cell = ws.cell(row=row, column=1, value='CONFIDENTIALITY NOTICE:')
                cell.font = Font(bold=True, size=9)
                row += 1
                ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
                ws.cell(row=row, column=1, value=notice_text)

        # Adjust column widths
        ws.column_dimensions['A'].width = 25
        ws.column_dimensions['B'].width = 40
        ws.column_dimensions['C'].width = 20
        ws.column_dimensions['D'].width = 20

        filepath = os.path.join(self.output_dir, filename)
        wb.save(filepath)
        return filepath

    def _add_xlsx_content(self, ws, row: int, doc_data: Dict[str, Any]) -> int:
        """Add content to worksheet."""
        skip_fields = {'document_id', 'document_type', 'category', 'subcategory',
                       'has_cui', 'classification', 'authority', 'distribution',
                       'generated_date', 'document_date', 'agency', 'title',
                       'confidentiality_notice'}

        for key, value in doc_data.items():
            if key in skip_fields:
                continue

            if isinstance(value, (str, int, float)):
                ws.cell(row=row, column=1, value=key.replace('_', ' ').title()).font = Font(bold=True)
                ws.cell(row=row, column=2, value=str(value))
                row += 1
            elif isinstance(value, list) and value:
                ws.cell(row=row, column=1, value=key.replace('_', ' ').title()).font = Font(bold=True)
                row += 1
                for item in value:
                    if isinstance(item, dict):
                        for k, v in item.items():
                            ws.cell(row=row, column=2, value=k)
                            ws.cell(row=row, column=3, value=str(v))
                            row += 1
                    else:
                        ws.cell(row=row, column=2, value=str(item))
                        row += 1
            elif isinstance(value, dict):
                ws.cell(row=row, column=1, value=key.replace('_', ' ').title()).font = Font(bold=True)
                row += 1
                for k, v in value.items():
                    ws.cell(row=row, column=2, value=k.replace('_', ' ').title())
                    ws.cell(row=row, column=3, value=str(v))
                    row += 1

        return row
